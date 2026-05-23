import io
import unittest

from openpyxl import Workbook, load_workbook

from app.standard_imports import (
    DOCUMENT_ATTRIBUTE_SHEET,
    DOCUMENT_TYPE_SHEET,
    META_SHEET,
    PBS_LEVEL_SHEET,
    STANDARD_SHEET,
    TAG_ATTRIBUTE_SHEET,
    TAG_CLASS_SHEET,
    build_standard_export_workbook_from_bundle,
    build_standard_import_template,
    validate_standard_import_workbook,
    validate_standard_import_rows,
    _extract_standard_import_rows,
)


def make_bundle() -> dict:
    return {
        "standard": {
            "id": "standard-1",
            "code": "DEC",
            "name": "DEC Engineering Standards",
            "version_label": "2024.1",
            "thumbnail_url": "data:image/webp;base64,abc",
            "status": "active",
            "metadata": {"source": "test"},
        },
        "pbs_levels": [
            {"level_no": 1, "code": "UNIT", "name": "装置", "description": "装置层"},
            {"level_no": 2, "code": "SYSTEM", "name": "系统", "description": None},
        ],
        "tag_classes": [
            {"code": "EQUIP", "name": "设备", "parent_code": None, "level_no": 1, "description": None, "status": "active"},
            {"code": "PUMP", "name": "泵", "parent_code": "EQUIP", "level_no": 2, "description": "泵设备", "status": "active"},
        ],
        "tag_attributes": [
            {
                "owner_class_code": "",
                "group_name": "公共",
                "code": "pressure",
                "name": "压力",
                "value_type": "number",
                "is_required": True,
                "unit_family": "pressure",
                "enum_options": [],
                "description": None,
                "sort_order": 0,
                "status": "active",
            },
            {
                "owner_class_code": "PUMP",
                "group_name": "",
                "code": "service",
                "name": "服务",
                "value_type": "enum",
                "is_required": True,
                "unit_family": None,
                "enum_options": ["FEED", "UTILITY"],
                "description": None,
                "sort_order": 1,
                "status": "active",
            },
        ],
        "document_types": [
            {
                "code": "DRAWING",
                "name": "图纸",
                "parent_code": None,
                "level_no": 1,
                "description": None,
                "status": "active",
                "allowed_extensions": [".dwg", ".pdf"],
                "metadata": {"review": "required"},
            }
        ],
        "document_attributes": [
            {
                "owner_document_type_code": "DRAWING",
                "group_name": "归档",
                "code": "discipline",
                "name": "专业",
                "value_type": "string",
                "is_required": True,
                "unit_family": None,
                "enum_options": [],
                "description": None,
                "sort_order": 0,
                "status": "active",
            }
        ],
    }


def make_minimal_workbook(*, metadata_json: str = "{}", duplicate_class: bool = False) -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = STANDARD_SHEET
    sheet.append(["code", "name", "version_label", "status", "thumbnail_url", "metadata_json"])
    sheet.append(["DEC", "DEC Engineering Standards", "2024.1", "active", "", metadata_json])

    pbs = workbook.create_sheet(PBS_LEVEL_SHEET)
    pbs.append(["level_no", "code", "name", "description"])
    pbs.append([1, "UNIT", "装置", ""])

    tag_classes = workbook.create_sheet(TAG_CLASS_SHEET)
    tag_classes.append(["code", "name", "parent_code", "level_no", "description", "status"])
    tag_classes.append(["PUMP", "泵", "", 1, "", "active"])
    if duplicate_class:
        tag_classes.append(["PUMP", "重复泵", "", 1, "", "active"])

    tag_attributes = workbook.create_sheet(TAG_ATTRIBUTE_SHEET)
    tag_attributes.append([
        "owner_class_code",
        "group_name",
        "code",
        "name",
        "value_type",
        "is_required",
        "unit_family",
        "enum_options",
        "description",
        "sort_order",
        "status",
    ])
    tag_attributes.append(["PUMP", "", "service", "服务", "enum", "true", "", "FEED,UTILITY", "", 0, "active"])

    document_types = workbook.create_sheet(DOCUMENT_TYPE_SHEET)
    document_types.append(["code", "name", "parent_code", "level_no", "description", "status", "allowed_extensions", "metadata_json"])
    document_types.append(["DRAWING", "图纸", "", 1, "", "active", ".dwg,.pdf", "{}"])

    document_attributes = workbook.create_sheet(DOCUMENT_ATTRIBUTE_SHEET)
    document_attributes.append([
        "owner_document_type_code",
        "group_name",
        "code",
        "name",
        "value_type",
        "is_required",
        "unit_family",
        "enum_options",
        "description",
        "sort_order",
        "status",
    ])

    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def make_generic_attribute_workbook() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "位号属性表"
    sheet.append(["属性编码", "属性名称", "数据类型", "是否必填", "说明"])
    sheet.append(["flow_rate", "流量", "数字", "是", "泵额定流量"])
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def make_cfihos_workbook() -> bytes:
    workbook = Workbook()
    document_type = workbook.active
    document_type.title = "document type"
    document_type.append([
        "CFIHOS unique code",
        "document type short code",
        "document type name",
        "document type description",
        "document type classification",
        "document type synonym name",
    ])
    document_type.append([
        "CFIHOS-70000004",
        "4301",
        "action list",
        "A collection of activities tracking responsible parties and status.",
        "DOCT",
        "punch list;issue list",
    ])

    discipline = workbook.create_sheet("discipline")
    discipline.append([
        "CFIHOS unique code",
        "discipline code",
        "discipline name",
        "discipline description",
    ])
    discipline.append([
        "CFIHOS-20000001",
        "ELE",
        "electrical",
        "Electrical engineering discipline.",
    ])

    discipline_document_type = workbook.create_sheet("discipline document type")
    discipline_document_type.append([
        "discipline document type CFIHOS unique code",
        "discipline CFIHOS unique code",
        "discipline code",
        "discipline name",
        "document type CFIHOS unique code",
        "document type short code",
        "document type name",
        "document type description",
        "discipline document type short code",
        "asset type reference",
        "representation type",
        "native file delivery timing",
    ])
    discipline_document_type.append([
        "CFIHOS-80000001",
        "CFIHOS-20000001",
        "ELE",
        "electrical",
        "CFIHOS-70000004",
        "4301",
        "action list",
        "A collection of activities tracking responsible parties and status.",
        "ELE-4301",
        "Tag",
        "Text",
        "during project",
    ])

    equipment_class = workbook.create_sheet("equipment class")
    equipment_class.append([
        "parent equipment class name",
        "equipment class CFIHOS unique code",
        "equipment class name",
        "equipment class definition",
        "abstract class indicator",
        "spare part information required indicator",
        "equipment class existence reason description",
        "equipment class synonym name",
    ])
    equipment_class.append([
        "",
        "CFIHOS-30000000",
        "mechanical equipment",
        "A parent equipment class.",
        "no",
        "yes",
        "",
        "",
    ])
    equipment_class.append([
        "mechanical equipment",
        "CFIHOS-30000001",
        "beam clamp",
        "An equipment item that can be clamped to a beam.",
        "no",
        "yes",
        "",
        "",
    ])

    equipment_property = workbook.create_sheet("equipment class property")
    equipment_property.append([
        "equipment class CFIHOS unique code",
        "equipment class name",
        "property CFIHOS unique code",
        "property name",
        "property relevant for equipment indicator",
        "property relevant for model / part indicator",
        "SI unit of measure CFIHOS unique code",
        "SI unit of measure name",
        "imperial unit of measure CFIHOS unique code",
        "imperial unit of measure name",
    ])
    equipment_property.append([
        "CFIHOS-30000001",
        "beam clamp",
        "CFIHOS-40000009",
        "ATEX category",
        "Yes",
        "",
        "",
        "",
        "",
        "",
    ])

    relationship = workbook.create_sheet("tag equipment class relationshi")
    relationship.append([
        "tag class CFIHOS unique code",
        "tag class name",
        "equipment class CFIHOS unique code",
        "equipment class name",
        "tag or equipment class relationship reason for mapping",
    ])
    relationship.append([
        "CFIHOS-10000001",
        "pressure transmitter",
        "CFIHOS-30000001",
        "beam clamp",
        "Test mapping",
    ])

    property_sheet = workbook.create_sheet("property")
    property_sheet.append([
        "CFIHOS unique code",
        "property name",
        "property definition",
        "property data type",
        "property data type length",
        "unit of measure dimension code CFIHOS unique code",
        "unit of measure dimension code",
        "property picklist name CFIHOS unique code",
        "property picklist name",
        "property existence reason description",
        "property synonym name",
    ])
    property_sheet.append([
        "CFIHOS-40000009",
        "ATEX category",
        "ATEX category definition",
        "Text",
        "",
        "",
        "",
        "CFIHOS-50000005",
        "ATEX category",
        "",
        "",
    ])

    picklist_sheet = workbook.create_sheet("property picklist values")
    picklist_sheet.append([
        "property picklist CFIHOS unique code",
        "property picklist name",
        "property picklist value CFIHOS unique code",
        "property picklist value code",
        "property picklist value description",
        "Source standard CFIHOS unique code",
        "source standard code",
    ])
    picklist_sheet.append(["CFIHOS-50000005", "ATEX category", "CFIHOS-60000006", "1", "Very high level", "", ""])
    picklist_sheet.append(["CFIHOS-50000005", "ATEX category", "CFIHOS-60000010", "2", "High level", "", ""])

    data_dictionary = workbook.create_sheet("data dictionary")
    data_dictionary.append([
        "CFIHOS unique code",
        "section",
        "object",
        "entity name",
        "property name",
        "entity attribute name",
        "definition",
        "note / comment",
        "example",
        "identifier / mandatory / optional",
        "format",
        "constraint must be present in",
        "data source",
        "former section (if different from last published version)",
        "relationship verb",
    ])
    data_dictionary.append([
        "CFIHOS-10000158",
        "",
        "attribute:",
        "equipment",
        "manufacturer company name",
        "",
        "A name used to uniquely identify the company who manufactures the equipment.",
        "",
        "Example : Flowserve",
        "Optional",
        "Text, max 30 characters",
        "model part",
        "",
        "",
        "is the model of",
    ])
    data_dictionary.append([
        "CFIHOS-10000163",
        "",
        "attribute:",
        "equipment",
        "equipment manufacturer serial number",
        "",
        "A unique identification number for the equipment as prescribed by the manufacturer.",
        "",
        "Example : AFKD34562837",
        "Optional",
        "Text, max 50 characters",
        "",
        "",
        "",
        "",
    ])
    data_dictionary.append([
        "CFIHOS-10000177",
        "",
        "attribute:",
        "tag",
        "tag description",
        "",
        "A functional description of the tag.",
        "",
        "",
        "Mandatory",
        "Text, max 255 characters",
        "",
        "",
        "",
        "",
    ])

    tag_class = workbook.create_sheet("tag class")
    tag_class.append([
        "parent tag class name",
        "CFIHOS unique code",
        "tag class name",
        "tag class definition",
        "abstract class indicator",
        "tag number format",
        "equipment expected to be installed indicator",
        "tag class existence reason description",
        "tag class synonym",
    ])
    tag_class.append([
        "measurement instrument",
        "CFIHOS-10000001",
        "pressure transmitter",
        "A device that measures pressure.",
        "no",
        "",
        "yes",
        "",
        "",
    ])

    document_required = workbook.create_sheet("document required per class")
    document_required.append([
        "source standard document and data requirement CFIHOS unique code",
        "tag or equipment class CFIHOS unique code",
        "tag or equipment class name",
        "asset type reference",
        "source standard CFIHOS unique code",
        "source standard code",
        "document type CFIHOS unique code",
        "document type name",
    ])
    document_required.append([
        "CFIHOS-90000001",
        "CFIHOS-10000001",
        "pressure transmitter",
        "Tag",
        "CFIHOS-90000177",
        "CFIHOS",
        "CFIHOS-70000004",
        "action list",
    ])
    document_required.append([
        "CFIHOS-90000002",
        "CFIHOS-10000001",
        "pressure transmitter",
        "Equipment",
        "CFIHOS-90000177",
        "CFIHOS",
        "CFIHOS-70000004",
        "action list",
    ])

    tag_property = workbook.create_sheet("tag class property")
    tag_property.append([
        "tag class CFIHOS unique code",
        "tag class name",
        "property CFIHOS unique code",
        "property name",
        "SI unit of measure CFIHOS unique code",
        "SI unit of measure name",
        "imperial unit of measure CFIHOS unique code",
        "imperial unit of measure name",
    ])
    tag_property.append([
        "CFIHOS-10000001",
        "pressure transmitter",
        "CFIHOS-40000010",
        "calibrated range",
        "CFIHOS-60000001",
        "bar",
        "",
        "",
    ])

    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


class StandardImportTemplateTest(unittest.TestCase):
    def test_builds_template_with_expected_sheets_and_headers(self):
        content = build_standard_import_template()
        workbook = load_workbook(io.BytesIO(content))

        self.assertIn("说明", workbook.sheetnames)
        self.assertIn(META_SHEET, workbook.sheetnames)
        self.assertIn(STANDARD_SHEET, workbook.sheetnames)
        self.assertIn(PBS_LEVEL_SHEET, workbook.sheetnames)
        self.assertIn(TAG_CLASS_SHEET, workbook.sheetnames)
        self.assertIn(TAG_ATTRIBUTE_SHEET, workbook.sheetnames)
        self.assertIn(DOCUMENT_TYPE_SHEET, workbook.sheetnames)
        self.assertIn(DOCUMENT_ATTRIBUTE_SHEET, workbook.sheetnames)
        self.assertEqual(workbook[META_SHEET].sheet_state, "hidden")
        self.assertEqual(
            [cell.value for cell in workbook[STANDARD_SHEET][1]],
            ["code", "name", "version_label", "status", "thumbnail_url", "metadata_json"],
        )

    def test_export_workbook_contains_complete_standard_definition(self):
        content = build_standard_export_workbook_from_bundle(make_bundle())
        workbook = load_workbook(io.BytesIO(content))

        self.assertEqual(workbook[STANDARD_SHEET]["A2"].value, "DEC")
        self.assertEqual(workbook[PBS_LEVEL_SHEET]["B3"].value, "SYSTEM")
        self.assertEqual(workbook[TAG_CLASS_SHEET]["A3"].value, "PUMP")
        self.assertEqual(workbook[TAG_ATTRIBUTE_SHEET]["C3"].value, "service")
        self.assertEqual(workbook[DOCUMENT_TYPE_SHEET]["A2"].value, "DRAWING")
        self.assertEqual(workbook[DOCUMENT_ATTRIBUTE_SHEET]["C2"].value, "discipline")


class StandardImportValidationTest(unittest.TestCase):
    def test_validates_new_standard_workbook_as_ready(self):
        result = validate_standard_import_workbook(
            make_minimal_workbook(),
            existing_standards=[],
        )

        self.assertEqual(result["summary"]["error_rows"], 0)
        self.assertEqual(result["summary"]["conflict_rows"], 0)
        self.assertTrue(result["summary"]["can_commit"])
        self.assertEqual(result["standard"]["code"], "DEC")
        self.assertEqual(len(result["rows"]), 5)

    def test_marks_existing_standard_code_as_conflict(self):
        result = validate_standard_import_workbook(
            make_minimal_workbook(),
            existing_standards=[{"id": "existing-standard", "code": "DEC"}],
        )

        self.assertEqual(result["summary"]["conflict_rows"], 1)
        self.assertFalse(result["summary"]["can_commit"])
        standard_row = next(row for row in result["rows"] if row["entity_kind"] == "standard")
        self.assertEqual(standard_row["status"], "conflict")
        self.assertEqual(standard_row["normalized_values"]["existing_standard_id"], "existing-standard")

    def test_reports_invalid_json_and_duplicate_codes(self):
        result = validate_standard_import_workbook(
            make_minimal_workbook(metadata_json="{bad-json", duplicate_class=True),
            existing_standards=[],
        )

        self.assertGreaterEqual(result["summary"]["error_rows"], 2)
        issue_codes = {
            issue["code"]
            for row in result["rows"]
            for issue in row["issues"]
        }
        self.assertIn("invalid_json", issue_codes)
        self.assertIn("duplicate_code_in_file", issue_codes)

    def test_rejects_missing_required_sheet(self):
        workbook = load_workbook(io.BytesIO(make_minimal_workbook()))
        del workbook[DOCUMENT_ATTRIBUTE_SHEET]
        buffer = io.BytesIO()
        workbook.save(buffer)

        with self.assertRaisesRegex(ValueError, "Workbook must contain"):
            validate_standard_import_workbook(buffer.getvalue(), existing_standards=[])

    def test_extracts_generic_xlsx_table_with_evidence(self):
        rows, chunks = _extract_standard_import_rows(
            "table-standard.xlsx",
            ".xlsx",
            make_generic_attribute_workbook(),
            target_mode="new",
            target_standard_bundle=None,
        )

        attribute_row = next(row for row in rows if row["entity_kind"] == "tag_attribute")
        self.assertEqual(attribute_row["values"]["code"], "flow_rate")
        self.assertEqual(attribute_row["values"]["name"], "流量")
        self.assertEqual(attribute_row["values"]["value_type"], "number")
        self.assertEqual(attribute_row["evidence"][0]["sheet_name"], "位号属性表")
        self.assertEqual(chunks[0]["source_kind"], "table")

    def test_extracts_cfihos_workbook_tables(self):
        rows, _ = _extract_standard_import_rows(
            "cfihos.xlsx",
            ".xlsx",
            make_cfihos_workbook(),
            target_mode="new",
            target_standard_bundle=None,
        )
        validated = validate_standard_import_rows(rows, existing_standards=[], target_mode="new")

        document_type = next(row for row in validated if row["entity_kind"] == "document_type")
        equipment_class = next(
            row
            for row in validated
            if row["entity_kind"] == "equipment_class" and row["normalized_values"]["code"] == "CFIHOS-30000001"
        )
        equipment_attribute = next(
            row
            for row in validated
            if row["entity_kind"] == "equipment_attribute"
            and row["normalized_values"]["owner_class_code"] == "CFIHOS-30000001"
        )
        equipment_common_attribute = next(
            row
            for row in validated
            if row["entity_kind"] == "equipment_attribute"
            and row["normalized_values"]["owner_class_code"] is None
            and row["normalized_values"]["code"] == "CFIHOS-10000158"
        )
        tag_attributes = [row for row in validated if row["entity_kind"] == "tag_attribute"]
        tag_attribute = next(
            row
            for row in tag_attributes
            if row["normalized_values"]["owner_class_code"] == "CFIHOS-10000001"
        )
        relationship = next(row for row in validated if row["entity_kind"] == "tag_equipment_class_relationship")
        discipline = next(row for row in validated if row["entity_kind"] == "discipline")
        discipline_document_type = next(row for row in validated if row["entity_kind"] == "discipline_document_type")
        class_document_requirements = [row for row in validated if row["entity_kind"] == "class_document_requirement"]
        class_document_requirement = class_document_requirements[0]

        self.assertEqual(document_type["normalized_values"]["code"], "4301")
        self.assertEqual(document_type["status"], "ready")
        self.assertEqual(discipline["normalized_values"]["code"], "ELE")
        self.assertEqual(discipline["normalized_values"]["cfihos_unique_code"], "CFIHOS-20000001")
        self.assertEqual(discipline["status"], "ready")
        self.assertEqual(discipline_document_type["normalized_values"]["discipline_code"], "ELE")
        self.assertEqual(discipline_document_type["normalized_values"]["document_type_code"], "4301")
        self.assertEqual(discipline_document_type["normalized_values"]["asset_scope"], "tag")
        self.assertEqual(discipline_document_type["normalized_values"]["lifecycle_phase"], "project")
        self.assertEqual(discipline_document_type["status"], "ready")
        self.assertEqual(class_document_requirement["normalized_values"]["class_code"], "CFIHOS-10000001")
        self.assertEqual(class_document_requirement["normalized_values"]["document_type_code"], "4301")
        self.assertEqual(class_document_requirement["normalized_values"]["source_standard_code"], "CFIHOS")
        self.assertEqual(class_document_requirement["status"], "ready")
        self.assertEqual({row["status"] for row in class_document_requirements}, {"ready"})
        self.assertEqual(equipment_class["normalized_values"]["name"], "beam clamp")
        self.assertEqual(equipment_class["normalized_values"]["parent_code"], "CFIHOS-30000000")
        self.assertEqual(equipment_class["normalized_values"]["level_no"], 2)
        self.assertEqual(equipment_class["status"], "ready")
        self.assertEqual(equipment_attribute["normalized_values"]["code"], "CFIHOS-40000009")
        self.assertEqual(equipment_attribute["normalized_values"]["name"], "ATEX category")
        self.assertEqual(equipment_attribute["normalized_values"]["value_type"], "enum")
        self.assertEqual(equipment_attribute["normalized_values"]["enum_options"], ["1", "2"])
        self.assertEqual(equipment_attribute["normalized_values"]["description"], "ATEX category definition")
        self.assertTrue(equipment_attribute["normalized_values"]["is_required"])
        self.assertEqual(equipment_attribute["status"], "ready")
        self.assertEqual(equipment_common_attribute["normalized_values"]["name"], "manufacturer company name")
        self.assertEqual(equipment_common_attribute["normalized_values"]["value_type"], "string")
        self.assertFalse(equipment_common_attribute["normalized_values"]["is_required"])
        self.assertEqual(equipment_common_attribute["status"], "ready")
        self.assertFalse(
            any(
                row["entity_kind"] == "tag_attribute" and row["normalized_values"]["code"] == "CFIHOS-10000177"
                for row in validated
            ),
            "CFIHOS data dictionary rows for tag must not be imported as equipment fixed attributes",
        )
        self.assertEqual(tag_attribute["normalized_values"]["code"], "CFIHOS-40000010")
        self.assertFalse(
            any(row["normalized_values"]["owner_class_code"] in (None, "") for row in tag_attributes),
            "CFIHOS property dictionary rows must not be imported as standard-level tag attributes",
        )
        self.assertEqual(relationship["normalized_values"]["tag_class_code"], "CFIHOS-10000001")
        self.assertEqual(relationship["normalized_values"]["equipment_class_code"], "CFIHOS-30000001")
        self.assertEqual(relationship["normalized_values"]["reason"], "Test mapping")
        self.assertEqual(relationship["status"], "ready")

    def test_reports_missing_cfihos_rule_references(self):
        workbook = load_workbook(io.BytesIO(make_cfihos_workbook()))
        workbook["discipline document type"]["E2"] = "CFIHOS-DOES-NOT-EXIST"
        workbook["discipline document type"]["F2"] = ""
        buffer = io.BytesIO()
        workbook.save(buffer)

        rows, _ = _extract_standard_import_rows(
            "cfihos.xlsx",
            ".xlsx",
            buffer.getvalue(),
            target_mode="new",
            target_standard_bundle=None,
        )
        validated = validate_standard_import_rows(rows, existing_standards=[], target_mode="new")
        rule = next(row for row in validated if row["entity_kind"] == "discipline_document_type")

        self.assertEqual(rule["status"], "error")
        self.assertIn("document_type_not_found", {issue["code"] for issue in rule["issues"]})

    def test_extracts_docx_tables_in_document_order(self):
        from docx import Document

        document = Document()
        document.add_heading("位号属性", level=1)
        table = document.add_table(rows=2, cols=4)
        headers = ["属性编码", "属性名称", "数据类型", "是否必填"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        values = ["pressure", "压力", "数字", "是"]
        for index, value in enumerate(values):
            table.rows[1].cells[index].text = value
        buffer = io.BytesIO()
        document.save(buffer)

        rows, chunks = _extract_standard_import_rows(
            "word-standard.docx",
            ".docx",
            buffer.getvalue(),
            target_mode="new",
            target_standard_bundle=None,
        )

        attribute_row = next(row for row in rows if row["entity_kind"] == "tag_attribute")
        self.assertEqual(attribute_row["values"]["code"], "pressure")
        self.assertEqual(attribute_row["values"]["name"], "压力")
        self.assertEqual(attribute_row["evidence"][0]["table_index"], 1)
        self.assertTrue(any(chunk["source_kind"] == "table" for chunk in chunks))

    def test_marks_merge_target_existing_items_as_conflicts(self):
        rows, _ = _extract_standard_import_rows(
            "table-standard.xlsx",
            ".xlsx",
            make_generic_attribute_workbook(),
            target_mode="merge",
            target_standard_bundle=make_bundle(),
        )
        rows.append(
            {
                "id": None,
                "row_number": len(rows) + 1,
                "source_kind": "table",
                "sheet_name": "位号类型",
                "page_no": None,
                "table_index": 1,
                "source_row_number": 2,
                "entity_kind": "tag_class",
                "values": {"code": "PUMP", "name": "泵", "parent_code": "", "level_no": 1, "status": "active"},
                "confidence": 0.8,
                "evidence": [],
            }
        )

        validated = validate_standard_import_rows(
            rows,
            existing_standards=[],
            target_mode="merge",
            target_standard_bundle=make_bundle(),
        )

        pump_row = next(row for row in validated if row["entity_kind"] == "tag_class" and row["normalized_values"]["code"] == "PUMP")
        self.assertEqual(pump_row["status"], "conflict")


if __name__ == "__main__":
    unittest.main()
