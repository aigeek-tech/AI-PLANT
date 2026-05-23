from __future__ import annotations

import io
import hashlib
import json
import math
import os
import re
from collections import Counter
from datetime import datetime, timezone
from typing import Any, Literal
from zipfile import BadZipFile

from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font
from openpyxl.utils.exceptions import InvalidFileException
from psycopg.types.json import Json

from .ai_client import complete_chat_text
from .db import fetch_all, fetch_one, get_connection
from .repository import get_ai_settings_secret


INSTRUCTION_SHEET = "说明"
META_SHEET = "_META"
STANDARD_SHEET = "标准"
PBS_LEVEL_SHEET = "PBS层级"
TAG_CLASS_SHEET = "位号类型"
TAG_ATTRIBUTE_SHEET = "位号属性"
EQUIPMENT_CLASS_SHEET = "设备类型"
EQUIPMENT_ATTRIBUTE_SHEET = "设备属性"
TAG_EQUIPMENT_CLASS_RELATIONSHIP_SHEET = "位号设备类型映射"
DOCUMENT_TYPE_SHEET = "文档类型"
DOCUMENT_ATTRIBUTE_SHEET = "文档属性"

SCHEMA_VERSION = "1"
DEFAULT_PAGE_SIZE = 50
MAX_PAGE_SIZE = 200
MAX_UPLOAD_SIZE = 10 * 1024 * 1024
SUPPORTED_UPLOAD_EXTENSIONS = {".xlsx", ".docx", ".pdf"}
TARGET_MODES = {"new", "merge"}
ITEM_ACTIONS = {"create", "update", "skip"}
AI_TABLE_PACKAGE_LIMIT = 24
MAX_AI_TABLE_ROWS = 80
TAG_EQUIPMENT_CLASS_RELATIONSHIP = "tag_equipment_class"

STANDARD_HEADERS = ["code", "name", "version_label", "status", "thumbnail_url", "metadata_json"]
PBS_LEVEL_HEADERS = ["level_no", "code", "name", "description"]
CLASS_HEADERS = ["code", "name", "parent_code", "level_no", "description", "status"]
TAG_ATTRIBUTE_HEADERS = [
    "owner_class_code",
    "group_name",
    "code",
    "name",
    "value_type",
    "is_required",
    "unit_family",
    "enum_options",
    "description",
    "sort_order",
    "status",
]
DOCUMENT_TYPE_HEADERS = ["code", "name", "parent_code", "level_no", "description", "status", "allowed_extensions", "metadata_json"]
DOCUMENT_ATTRIBUTE_HEADERS = [
    "owner_document_type_code",
    "group_name",
    "code",
    "name",
    "value_type",
    "is_required",
    "unit_family",
    "enum_options",
    "description",
    "sort_order",
    "status",
]
TAG_EQUIPMENT_CLASS_RELATIONSHIP_HEADERS = ["tag_class_code", "equipment_class_code", "reason", "status"]

REQUIRED_SHEETS = {
    STANDARD_SHEET: STANDARD_HEADERS,
    PBS_LEVEL_SHEET: PBS_LEVEL_HEADERS,
    TAG_CLASS_SHEET: CLASS_HEADERS,
    TAG_ATTRIBUTE_SHEET: TAG_ATTRIBUTE_HEADERS,
    DOCUMENT_TYPE_SHEET: DOCUMENT_TYPE_HEADERS,
    DOCUMENT_ATTRIBUTE_SHEET: DOCUMENT_ATTRIBUTE_HEADERS,
}
VALID_STANDARD_STATUSES = {"draft", "active", "archived"}
VALID_DEFINITION_STATUSES = {"draft", "active", "deprecated", "archived"}
VALID_RULE_STATUSES = {"active", "deprecated", "archived"}
VALID_VALUE_TYPES = {"string", "number", "integer", "boolean", "date", "enum", "json"}
CONFLICT_ACTIONS = {"create_copy", "merge_update", "skip"}
ROW_STATUS_VALUES = {"ready", "error", "warning", "conflict"}

HEADER_ALIASES = {
    "code": {"code", "编码", "编号", "代号", "类型编码", "类别编码", "class_code", "class code", "document_type_code"},
    "name": {"name", "名称", "类型名称", "类别名称", "属性名称", "中文名称", "标准名称"},
    "standard_code": {"standard_code", "标准编码", "标准编号", "标准号"},
    "version_label": {"version_label", "版本", "版本号", "版次", "version"},
    "parent_code": {"parent_code", "父级编码", "上级编码", "父类编码", "parent"},
    "level_no": {"level_no", "层级", "层级号", "级别", "序号", "level"},
    "description": {"description", "说明", "描述", "备注", "释义"},
    "status": {"status", "状态"},
    "owner_class_code": {"owner_class_code", "所属类别", "类别编码", "位号类型", "class", "class_code"},
    "owner_document_type_code": {"owner_document_type_code", "所属文档类型", "文档类型", "文件类型"},
    "value_type": {"value_type", "数据类型", "类型", "字段类型", "属性类型"},
    "is_required": {"is_required", "是否必填", "必填", "是否必须", "required"},
    "unit_family": {"unit_family", "单位", "单位族", "量纲"},
    "enum_options": {"enum_options", "枚举", "枚举值", "可选值", "选项"},
    "sort_order": {"sort_order", "排序", "排序号"},
    "allowed_extensions": {"allowed_extensions", "扩展名", "文件扩展名", "允许格式", "格式"},
}

KNOWN_SOURCE_TABLE_KINDS = {
    "document type": "document_type",
    "discipline": "discipline",
    "discipline document type": "discipline_document_type",
    "equipment class": "equipment_class",
    "tag class": "tag_class",
    "data dictionary": "equipment_attribute",
    "equipment class property": "equipment_attribute",
    "tag class property": "tag_attribute",
    "tag equipment class relationshi": "tag_equipment_class_relationship",
    "document required per class": "class_document_requirement",
}

UNSUPPORTED_SOURCE_TABLES = {
    "cover and index",
    "rdl master object",
    "cfihos object equivalent mappin",
    "handover event",
    "property picklist values",
    "property groupings",
    "source standard",
    "jip33 info required spec",
    "tag or equip class src standard",
    "tag equip class prop src std",
    "unit of measure",
}


def build_standard_import_template() -> bytes:
    workbook = Workbook()
    instruction_sheet = workbook.active
    instruction_sheet.title = INSTRUCTION_SHEET
    instructions = [
        "1. 仅支持 .xlsx，上传后会先生成导入预览，不会立即写入标准库。",
        "2. 每个文件只描述一个标准；标准 Sheet 必须且只能填写一行有效标准信息。",
        "3. 位号属性 owner_class_code 留空表示标准公共属性；文档属性 owner_document_type_code 留空表示公共文档属性。",
        "4. enum_options、allowed_extensions 可用英文逗号或分号分隔；metadata_json 必须是 JSON 对象。",
        "5. 同编码标准会进入冲突预览，可选择创建副本、非破坏合并更新或跳过。",
        "6. 非破坏合并只新增或更新文件中出现的定义，不会删除目标库里额外存在的定义。",
    ]
    for row, line in enumerate(instructions, start=1):
        instruction_sheet.cell(row=row, column=1, value=line)
    instruction_sheet.column_dimensions["A"].width = 120

    _write_meta_sheet(workbook, source_standard=None)
    _create_sheet(workbook, STANDARD_SHEET, STANDARD_HEADERS, [["DEC", "DEC Engineering Standards", "2024.1", "active", "", "{}"]])
    _create_sheet(workbook, PBS_LEVEL_SHEET, PBS_LEVEL_HEADERS, [[1, "UNIT", "装置", ""]])
    _create_sheet(workbook, TAG_CLASS_SHEET, CLASS_HEADERS, [["PUMP", "泵", "", 1, "", "active"]])
    _create_sheet(
        workbook,
        TAG_ATTRIBUTE_SHEET,
        TAG_ATTRIBUTE_HEADERS,
        [["PUMP", "", "service", "服务", "enum", "true", "", "FEED,UTILITY", "", 0, "active"]],
    )
    _create_sheet(workbook, EQUIPMENT_CLASS_SHEET, CLASS_HEADERS, [["CENTRIFUGAL_PUMP", "离心泵", "", 1, "", "active"]])
    _create_sheet(
        workbook,
        EQUIPMENT_ATTRIBUTE_SHEET,
        TAG_ATTRIBUTE_HEADERS,
        [["CENTRIFUGAL_PUMP", "", "serial_no", "序列号", "string", "false", "", "", "", 0, "active"]],
    )
    _create_sheet(
        workbook,
        TAG_EQUIPMENT_CLASS_RELATIONSHIP_SHEET,
        TAG_EQUIPMENT_CLASS_RELATIONSHIP_HEADERS,
        [["PUMP", "CENTRIFUGAL_PUMP", "泵位号可由离心泵实物实现", "active"]],
    )
    _create_sheet(
        workbook,
        DOCUMENT_TYPE_SHEET,
        DOCUMENT_TYPE_HEADERS,
        [["DRAWING", "图纸", "", 1, "", "active", ".dwg,.pdf", "{}"]],
    )
    _create_sheet(workbook, DOCUMENT_ATTRIBUTE_SHEET, DOCUMENT_ATTRIBUTE_HEADERS, [])

    return _save_workbook(workbook)


def build_standard_export_workbook(standard_id: str) -> tuple[dict, bytes] | None:
    bundle = load_standard_export_bundle(standard_id)
    if bundle is None:
        return None
    return bundle["standard"], build_standard_export_workbook_from_bundle(bundle)


def build_standard_export_workbook_from_bundle(bundle: dict) -> bytes:
    workbook = Workbook()
    instruction_sheet = workbook.active
    instruction_sheet.title = INSTRUCTION_SHEET
    for row, line in enumerate(
        [
            f"导出标准: {bundle['standard']['code']} {bundle['standard']['name']}",
            "1. 本文件可直接重新上传到标准库导入入口。",
            "2. 导出内容仅包含标准定义，不包含项目、TAG 实例、图纸实例、权限或项目绑定。",
            "3. 重新导入到同编码标准时，系统会先预览冲突并要求选择处理方式。",
        ],
        start=1,
    ):
        instruction_sheet.cell(row=row, column=1, value=line)
    instruction_sheet.column_dimensions["A"].width = 120

    _write_meta_sheet(workbook, source_standard=bundle["standard"])
    standard = bundle["standard"]
    _create_sheet(
        workbook,
        STANDARD_SHEET,
        STANDARD_HEADERS,
        [[
            standard["code"],
            standard["name"],
            standard.get("version_label"),
            standard.get("status", "active"),
            standard.get("thumbnail_url"),
            _json_text(standard.get("metadata") or {}),
        ]],
    )
    _create_sheet(
        workbook,
        PBS_LEVEL_SHEET,
        PBS_LEVEL_HEADERS,
        [
            [row.get("level_no"), row.get("code"), row.get("name"), row.get("description")]
            for row in bundle.get("pbs_levels", [])
        ],
    )
    _create_sheet(
        workbook,
        TAG_CLASS_SHEET,
        CLASS_HEADERS,
        [
            [row.get("code"), row.get("name"), row.get("parent_code"), row.get("level_no"), row.get("description"), row.get("status")]
            for row in bundle.get("tag_classes", [])
        ],
    )
    _create_sheet(
        workbook,
        TAG_ATTRIBUTE_SHEET,
        TAG_ATTRIBUTE_HEADERS,
        [
            [
                row.get("owner_class_code") or "",
                row.get("group_name") or "",
                row.get("code"),
                row.get("name"),
                row.get("value_type"),
                "true" if row.get("is_required") else "false",
                row.get("unit_family") or "",
                ",".join(map(str, row.get("enum_options") or [])),
                row.get("description"),
                row.get("sort_order"),
                row.get("status"),
            ]
            for row in bundle.get("tag_attributes", [])
        ],
    )
    _create_sheet(
        workbook,
        EQUIPMENT_CLASS_SHEET,
        CLASS_HEADERS,
        [
            [row.get("code"), row.get("name"), row.get("parent_code"), row.get("level_no"), row.get("description"), row.get("status")]
            for row in bundle.get("equipment_classes", [])
        ],
    )
    _create_sheet(
        workbook,
        EQUIPMENT_ATTRIBUTE_SHEET,
        TAG_ATTRIBUTE_HEADERS,
        [
            [
                row.get("owner_class_code") or "",
                row.get("group_name") or "",
                row.get("code"),
                row.get("name"),
                row.get("value_type"),
                "true" if row.get("is_required") else "false",
                row.get("unit_family") or "",
                ",".join(map(str, row.get("enum_options") or [])),
                row.get("description"),
                row.get("sort_order"),
                row.get("status"),
            ]
            for row in bundle.get("equipment_attributes", [])
        ],
    )
    _create_sheet(
        workbook,
        TAG_EQUIPMENT_CLASS_RELATIONSHIP_SHEET,
        TAG_EQUIPMENT_CLASS_RELATIONSHIP_HEADERS,
        [
            [row.get("tag_class_code"), row.get("equipment_class_code"), row.get("reason"), row.get("status")]
            for row in bundle.get("tag_equipment_class_relationships", [])
        ],
    )
    _create_sheet(
        workbook,
        DOCUMENT_TYPE_SHEET,
        DOCUMENT_TYPE_HEADERS,
        [
            [
                row.get("code"),
                row.get("name"),
                row.get("parent_code"),
                row.get("level_no"),
                row.get("description"),
                row.get("status"),
                ",".join(map(str, row.get("allowed_extensions") or [])),
                _json_text(row.get("metadata") or {}),
            ]
            for row in bundle.get("document_types", [])
        ],
    )
    _create_sheet(
        workbook,
        DOCUMENT_ATTRIBUTE_SHEET,
        DOCUMENT_ATTRIBUTE_HEADERS,
        [
            [
                row.get("owner_document_type_code") or "",
                row.get("group_name") or "",
                row.get("code"),
                row.get("name"),
                row.get("value_type"),
                "true" if row.get("is_required") else "false",
                row.get("unit_family") or "",
                ",".join(map(str, row.get("enum_options") or [])),
                row.get("description"),
                row.get("sort_order"),
                row.get("status"),
            ]
            for row in bundle.get("document_attributes", [])
        ],
    )
    return _save_workbook(workbook)


def load_standard_export_bundle(standard_id: str) -> dict | None:
    standard = fetch_one(
        """
        SELECT id, code, name, version_label, thumbnail_url, status, metadata
        FROM standard
        WHERE id = %s
        """,
        (standard_id,),
    )
    if standard is None:
        return None

    tag_classes = fetch_all(
        """
        SELECT c.id, c.code, c.name, parent.code AS parent_code, c.level_no, c.description, c.status
        FROM class c
        LEFT JOIN class parent ON parent.id = c.parent_id
        WHERE c.standard_id = %s
          AND c.applies_to IN ('tag', 'both')
        ORDER BY c.level_no, c.code
        """,
        (standard_id,),
    )
    tag_attributes = fetch_all(
        """
        SELECT
            owner.code AS owner_class_code,
            ad.group_name,
            ad.code,
            ad.name,
            ad.value_type,
            ad.is_required,
            ad.unit_family,
            ad.enum_options,
            ad.description,
            ad.sort_order,
            ad.status
        FROM attribute_definition ad
        LEFT JOIN class owner ON owner.id = ad.class_id
        WHERE (owner.standard_id = %s OR ad.standard_id = %s)
          AND ad.applies_to IN ('tag', 'both')
          AND ad.status <> 'archived'
        ORDER BY owner.code NULLS FIRST, ad.sort_order, ad.code
        """,
        (standard_id, standard_id),
    )
    pbs_levels = fetch_all(
        """
        SELECT level_no, code, name, description
        FROM pbs_level_template
        WHERE standard_id = %s
        ORDER BY level_no, code
        """,
        (standard_id,),
    )
    document_types = fetch_all(
        """
        SELECT
            c.id,
            c.code,
            c.name,
            parent.code AS parent_code,
            c.level_no,
            c.description,
            c.status,
            COALESCE(c.metadata -> 'document' -> 'allowed_extensions', '[]'::jsonb) AS allowed_extensions,
            c.metadata
        FROM class c
        LEFT JOIN class parent ON parent.id = c.parent_id
        WHERE c.standard_id = %s
          AND c.applies_to IN ('document', 'both')
        ORDER BY c.level_no, c.code
        """,
        (standard_id,),
    )
    document_attributes = fetch_all(
        """
        SELECT
            owner.code AS owner_document_type_code,
            ad.group_name,
            ad.code,
            ad.name,
            ad.value_type,
            ad.is_required,
            ad.unit_family,
            ad.enum_options,
            ad.description,
            ad.sort_order,
            ad.status
        FROM attribute_definition ad
        LEFT JOIN class owner ON owner.id = ad.class_id
        WHERE (owner.standard_id = %s OR ad.standard_id = %s)
          AND ad.applies_to IN ('document', 'both')
          AND ad.status <> 'archived'
        ORDER BY owner.code NULLS FIRST, ad.sort_order, ad.code
        """,
        (standard_id, standard_id),
    )
    equipment_classes = fetch_all(
        """
        SELECT c.id, c.code, c.name, parent.code AS parent_code, c.level_no, c.description, c.status
        FROM class c
        LEFT JOIN class parent ON parent.id = c.parent_id
        WHERE c.standard_id = %s
          AND c.applies_to = 'equipment'
        ORDER BY c.level_no, c.code
        """,
        (standard_id,),
    )
    equipment_attributes = fetch_all(
        """
        SELECT
            owner.code AS owner_class_code,
            ad.group_name,
            ad.code,
            ad.name,
            ad.value_type,
            ad.is_required,
            ad.unit_family,
            ad.enum_options,
            ad.description,
            ad.sort_order,
            ad.status
        FROM attribute_definition ad
        LEFT JOIN class owner ON owner.id = ad.class_id
        WHERE (owner.standard_id = %s OR ad.standard_id = %s)
          AND ad.applies_to = 'equipment'
          AND ad.status <> 'archived'
        ORDER BY owner.code NULLS FIRST, ad.sort_order, ad.code
        """,
        (standard_id, standard_id),
    )
    tag_equipment_class_relationships = fetch_all(
        """
        SELECT
            source.code AS tag_class_code,
            target.code AS equipment_class_code,
            cr.reason,
            cr.status
        FROM class_relationship cr
        JOIN class source ON source.id = cr.source_class_id
        JOIN class target ON target.id = cr.target_class_id
        WHERE cr.standard_id = %s
          AND cr.relationship_type = %s
          AND source.applies_to IN ('tag', 'both')
          AND target.applies_to = 'equipment'
          AND cr.status <> 'archived'
        ORDER BY source.code, target.code
        """,
        (standard_id, TAG_EQUIPMENT_CLASS_RELATIONSHIP),
    )

    return {
        "standard": standard,
        "pbs_levels": pbs_levels,
        "tag_classes": tag_classes,
        "tag_attributes": tag_attributes,
        "equipment_classes": equipment_classes,
        "equipment_attributes": equipment_attributes,
        "tag_equipment_class_relationships": tag_equipment_class_relationships,
        "document_types": document_types,
        "document_attributes": document_attributes,
    }


def validate_standard_import_workbook(
    workbook_bytes: bytes,
    *,
    existing_standards: list[dict] | None = None,
    target_mode: str = "new",
    target_standard_bundle: dict | None = None,
) -> dict:
    try:
        row_values = _parse_standard_import_workbook(workbook_bytes)
    except (BadZipFile, InvalidFileException, OSError, ValueError) as error:
        message = str(error)
        if isinstance(error, (BadZipFile, InvalidFileException, OSError)):
            message = "Uploaded file is not a valid .xlsx workbook"
        raise ValueError(message) from error

    rows = validate_standard_import_rows(
        row_values,
        existing_standards=existing_standards or [],
        target_mode=target_mode,
        target_standard_bundle=target_standard_bundle,
    )
    standard_row = next((row for row in rows if row["entity_kind"] == "standard"), None)
    return {
        "standard": standard_row["normalized_values"] if standard_row else {},
        "summary": _summarize_rows(rows),
        "rows": rows,
    }


def validate_standard_import_rows(
    row_values: list[dict],
    *,
    existing_standards: list[dict],
    target_mode: str = "new",
    target_standard_bundle: dict | None = None,
) -> list[dict]:
    existing_by_code = {
        _code_key(row.get("code")): row
        for row in existing_standards
        if _code_key(row.get("code"))
    }
    by_kind: dict[str, list[dict]] = {}
    for row in row_values:
        by_kind.setdefault(row["entity_kind"], []).append(row)

    tag_class_codes = [_normalize_code(row.get("values", {}).get("code")) for row in by_kind.get("tag_class", [])]
    equipment_class_codes = [_normalize_code(row.get("values", {}).get("code")) for row in by_kind.get("equipment_class", [])]
    document_type_codes = [_normalize_code(row.get("values", {}).get("code")) for row in by_kind.get("document_type", [])]
    discipline_codes = [_normalize_code(row.get("values", {}).get("code")) for row in by_kind.get("discipline", [])]
    discipline_cfihos_codes = [_normalize_code(row.get("values", {}).get("cfihos_unique_code")) for row in by_kind.get("discipline", [])]
    document_type_cfihos_codes: list[str] = []
    document_type_code_by_cfihos: dict[str, str] = {}
    for row in by_kind.get("document_type", []):
        metadata, _ = _parse_json_object(row.get("values", {}).get("metadata_json"), "metadata_json")
        cfihos_code = _normalize_code(metadata.get("cfihos_unique_code"))
        document_type_code = _normalize_code(row.get("values", {}).get("code"))
        if cfihos_code:
            document_type_cfihos_codes.append(cfihos_code)
            if document_type_code:
                document_type_code_by_cfihos[_code_key(cfihos_code)] = document_type_code
    tag_class_levels = _infer_class_levels(by_kind.get("tag_class", []))
    equipment_class_levels = _infer_class_levels(by_kind.get("equipment_class", []))
    document_type_levels = _infer_class_levels(by_kind.get("document_type", []))
    context = {
        "pbs_code_counts": Counter(_code_key(row.get("values", {}).get("code")) for row in by_kind.get("pbs_level", [])),
        "pbs_level_counts": Counter(_int_or_none(row.get("values", {}).get("level_no")) for row in by_kind.get("pbs_level", [])),
        "tag_class_code_counts": Counter(_code_key(code) for code in tag_class_codes),
        "equipment_class_code_counts": Counter(_code_key(code) for code in equipment_class_codes),
        "document_type_code_counts": Counter(_code_key(code) for code in document_type_codes),
        "discipline_code_counts": Counter(_code_key(code) for code in discipline_codes),
        "discipline_cfihos_code_counts": Counter(_code_key(code) for code in discipline_cfihos_codes if _code_key(code)),
        "tag_class_codes": {_code_key(code) for code in tag_class_codes if _code_key(code)},
        "equipment_class_codes": {_code_key(code) for code in equipment_class_codes if _code_key(code)},
        "document_type_codes": {_code_key(code) for code in document_type_codes if _code_key(code)},
        "document_type_cfihos_codes": {_code_key(code) for code in document_type_cfihos_codes if _code_key(code)},
        "document_type_code_by_cfihos": document_type_code_by_cfihos,
        "discipline_codes": {_code_key(code) for code in discipline_codes if _code_key(code)},
        "discipline_cfihos_codes": {_code_key(code) for code in discipline_cfihos_codes if _code_key(code)},
        "tag_class_levels": tag_class_levels,
        "equipment_class_levels": equipment_class_levels,
        "document_type_levels": document_type_levels,
        "target_mode": target_mode,
        "target_existing": _build_target_existing_index(target_standard_bundle),
        "tag_attribute_counts": Counter(
            (_code_key(row.get("values", {}).get("owner_class_code")) or "__common__", _code_key(row.get("values", {}).get("code")))
            for row in by_kind.get("tag_attribute", [])
        ),
        "equipment_attribute_counts": Counter(
            (_code_key(row.get("values", {}).get("owner_class_code")) or "__common__", _code_key(row.get("values", {}).get("code")))
            for row in by_kind.get("equipment_attribute", [])
        ),
        "document_attribute_counts": Counter(
            (_code_key(row.get("values", {}).get("owner_document_type_code")) or "__common__", _code_key(row.get("values", {}).get("code")))
            for row in by_kind.get("document_attribute", [])
        ),
        "tag_equipment_relationship_counts": Counter(
            (
                _code_key(row.get("values", {}).get("tag_class_code")),
                _code_key(row.get("values", {}).get("equipment_class_code")),
            )
            for row in by_kind.get("tag_equipment_class_relationship", [])
        ),
        "discipline_document_type_counts": Counter(
            _discipline_document_type_key(row.get("values", {}))
            for row in by_kind.get("discipline_document_type", [])
        ),
        "class_document_requirement_counts": Counter(
            _class_document_requirement_key(row.get("values", {}))
            for row in by_kind.get("class_document_requirement", [])
        ),
    }

    validated_rows: list[dict] = []
    for row in row_values:
        kind = row["entity_kind"]
        if kind == "standard":
            validated_rows.append(_validate_standard_row(row, existing_by_code))
        elif kind == "pbs_level":
            validated_rows.append(_validate_pbs_level_row(row, context))
        elif kind == "tag_class":
            validated_rows.append(_validate_class_row(row, context, kind="tag_class"))
        elif kind == "equipment_class":
            validated_rows.append(_validate_class_row(row, context, kind="equipment_class"))
        elif kind == "tag_attribute":
            validated_rows.append(_validate_attribute_row(row, context, kind="tag_attribute"))
        elif kind == "equipment_attribute":
            validated_rows.append(_validate_attribute_row(row, context, kind="equipment_attribute"))
        elif kind == "document_type":
            validated_rows.append(_validate_document_type_row(row, context))
        elif kind == "document_attribute":
            validated_rows.append(_validate_attribute_row(row, context, kind="document_attribute"))
        elif kind == "tag_equipment_class_relationship":
            validated_rows.append(_validate_tag_equipment_relationship_row(row, context))
        elif kind == "discipline":
            validated_rows.append(_validate_discipline_row(row, context))
        elif kind == "discipline_document_type":
            validated_rows.append(_validate_discipline_document_type_row(row, context))
        elif kind == "class_document_requirement":
            validated_rows.append(_validate_class_document_requirement_row(row, context))
    return validated_rows


def _infer_class_levels(rows: list[dict]) -> dict[str, int]:
    by_code = {
        _code_key(row.get("values", {}).get("code")): row
        for row in rows
        if _code_key(row.get("values", {}).get("code"))
    }
    levels: dict[str, int] = {}

    def resolve(code_key: str, visiting: set[str]) -> int:
        if code_key in levels:
            return levels[code_key]
        if code_key in visiting:
            return 1
        row = by_code.get(code_key)
        if row is None:
            return 1
        explicit = _int_or_none(row.get("values", {}).get("level_no"))
        if explicit and explicit > 0:
            levels[code_key] = explicit
            return explicit
        parent_key = _code_key(row.get("values", {}).get("parent_code"))
        if parent_key and parent_key in by_code:
            levels[code_key] = resolve(parent_key, {*visiting, code_key}) + 1
        else:
            levels[code_key] = 1
        return levels[code_key]

    for key in by_code:
        resolve(key, set())
    return levels


def _relationship_text_key(value: Any) -> str:
    return _trim_text(value).lower()


def _discipline_document_type_key(values: dict) -> tuple[str, str, str, str, str, str, str]:
    discipline_ref = _code_key(
        values.get("discipline_cfihos_unique_code")
        or values.get("discipline_code")
        or values.get("discipline_id")
    )
    document_type_ref = _code_key(
        values.get("document_type_cfihos_unique_code")
        or values.get("document_type_code")
        or values.get("document_type_id")
    )
    return (
        discipline_ref,
        document_type_ref,
        _normalize_context_code(values.get("asset_scope"), default="") or "",
        _relationship_text_key(values.get("representation_type")),
        _relationship_text_key(values.get("native_file_delivery_timing")),
        _normalize_context_code(values.get("perspective"), default="standard") or "standard",
        _normalize_context_code(values.get("lifecycle_phase"), default="unspecified") or "unspecified",
    )


def _class_document_requirement_key(values: dict) -> tuple[str, str, str, str, str, str]:
    class_ref = _code_key(values.get("class_code") or values.get("class_id"))
    document_type_ref = _code_key(
        values.get("document_type_cfihos_unique_code")
        or values.get("document_type_code")
        or values.get("document_type_id")
    )
    return (
        class_ref,
        document_type_ref,
        _normalize_context_code(values.get("asset_scope"), default="") or "",
        _code_key(values.get("source_standard_code") or values.get("source_standard_cfihos_code")),
        _normalize_context_code(values.get("perspective"), default="standard") or "standard",
        _normalize_context_code(values.get("lifecycle_phase"), default="unspecified") or "unspecified",
    )


def _build_target_existing_index(bundle: dict | None) -> dict[str, set[tuple[str, str] | str]]:
    if not bundle:
        return {
            "pbs_level": set(),
            "tag_class": set(),
            "equipment_class": set(),
            "document_type": set(),
            "tag_attribute": set(),
            "equipment_attribute": set(),
            "document_attribute": set(),
            "tag_equipment_class_relationship": set(),
            "discipline": set(),
            "discipline_cfihos": set(),
            "discipline_document_type": set(),
            "class_document_requirement": set(),
        }
    return {
        "pbs_level": {_code_key(row.get("code")) for row in bundle.get("pbs_levels", []) if _code_key(row.get("code"))},
        "tag_class": {_code_key(row.get("code")) for row in bundle.get("tag_classes", []) if _code_key(row.get("code"))},
        "equipment_class": {_code_key(row.get("code")) for row in bundle.get("equipment_classes", []) if _code_key(row.get("code"))},
        "document_type": {_code_key(row.get("code")) for row in bundle.get("document_types", []) if _code_key(row.get("code"))},
        "tag_attribute": {
            (_code_key(row.get("owner_class_code")) or "__common__", _code_key(row.get("code")))
            for row in bundle.get("tag_attributes", [])
            if _code_key(row.get("code"))
        },
        "equipment_attribute": {
            (_code_key(row.get("owner_class_code")) or "__common__", _code_key(row.get("code")))
            for row in bundle.get("equipment_attributes", [])
            if _code_key(row.get("code"))
        },
        "document_attribute": {
            (_code_key(row.get("owner_document_type_code")) or "__common__", _code_key(row.get("code")))
            for row in bundle.get("document_attributes", [])
            if _code_key(row.get("code"))
        },
        "tag_equipment_class_relationship": {
            (_code_key(row.get("tag_class_code")), _code_key(row.get("equipment_class_code")))
            for row in bundle.get("tag_equipment_class_relationships", [])
            if _code_key(row.get("tag_class_code")) and _code_key(row.get("equipment_class_code"))
        },
        "discipline": {_code_key(row.get("code")) for row in bundle.get("disciplines", []) if _code_key(row.get("code"))},
        "discipline_cfihos": {_code_key(row.get("cfihos_unique_code")) for row in bundle.get("disciplines", []) if _code_key(row.get("cfihos_unique_code"))},
        "discipline_document_type": {
            _discipline_document_type_key(row)
            for row in bundle.get("discipline_document_types", [])
        },
        "class_document_requirement": {
            _class_document_requirement_key(row)
            for row in bundle.get("class_document_requirements", [])
        },
    }


def _committable_values(rows: list[dict], entity_kind: str) -> list[dict]:
    return [
        row["normalized_values"]
        for row in rows
        if row["entity_kind"] == entity_kind and row.get("action") != "skip"
    ]


def create_standard_import_job_from_upload(
    filename: str,
    file_bytes: bytes,
    *,
    target_mode: str = "new",
    target_standard_id: str | None = None,
) -> dict:
    normalized_filename = os.path.basename(filename.strip())
    if not normalized_filename:
        raise ValueError("Uploaded file must have a filename")

    file_ext = os.path.splitext(normalized_filename)[1].lower()
    if file_ext not in SUPPORTED_UPLOAD_EXTENSIONS:
        raise ValueError("Only .docx, .xlsx, and text-based .pdf files are supported")
    if len(file_bytes) > MAX_UPLOAD_SIZE:
        raise ValueError("Uploaded file exceeds the 10 MB limit")

    normalized_target_mode = target_mode if target_mode in TARGET_MODES else "new"
    target_bundle = None
    if normalized_target_mode == "merge":
        if not target_standard_id:
            raise ValueError("target_standard_id is required when target_mode is merge")
        target_bundle = load_standard_export_bundle(target_standard_id)
        if target_bundle is None:
            raise ValueError("Target standard not found")
    else:
        target_standard_id = None

    existing_standards = fetch_all("SELECT id, code FROM standard")
    row_values, chunks = _extract_standard_import_rows(
        normalized_filename,
        file_ext,
        file_bytes,
        target_mode=normalized_target_mode,
        target_standard_bundle=target_bundle,
    )
    validated = _validate_extracted_standard_rows(
        row_values,
        existing_standards=existing_standards,
        target_mode=normalized_target_mode,
        target_standard_bundle=target_bundle,
    )
    created = _store_standard_import_job(
        filename=normalized_filename,
        file_ext=file_ext.lstrip("."),
        file_size=len(file_bytes),
        checksum_sha256=hashlib.sha256(file_bytes).hexdigest(),
        target_mode=normalized_target_mode,
        target_standard_id=target_standard_id,
        validated=validated,
        chunks=chunks,
    )
    return get_standard_import_job_detail(str(created["id"]))


def _extract_standard_import_rows(
    filename: str,
    file_ext: str,
    file_bytes: bytes,
    *,
    target_mode: str,
    target_standard_bundle: dict | None,
) -> tuple[list[dict], list[dict]]:
    if file_ext == ".xlsx" and _is_standard_template_workbook(file_bytes):
        return _parse_standard_import_workbook(file_bytes), []

    if file_ext == ".xlsx":
        packages, chunks = _extract_xlsx_table_packages(filename, file_bytes)
    elif file_ext == ".docx":
        packages, chunks = _extract_docx_table_packages(filename, file_bytes)
    else:
        packages, chunks = _extract_pdf_context_packages(filename, file_bytes)

    ai_rows = _build_ai_candidate_rows(filename, packages, target_mode=target_mode)
    rule_rows = _packages_to_import_rows(filename, packages)
    rows = _merge_candidate_rows([*ai_rows, *rule_rows])
    if not any(row["entity_kind"] == "standard" for row in rows):
        rows.insert(0, _fallback_standard_row(filename, file_ext, target_mode, target_standard_bundle))

    return _renumber_import_rows(rows), chunks


def _validate_extracted_standard_rows(
    row_values: list[dict],
    *,
    existing_standards: list[dict],
    target_mode: str,
    target_standard_bundle: dict | None,
) -> dict:
    rows = validate_standard_import_rows(
        row_values,
        existing_standards=existing_standards,
        target_mode=target_mode,
        target_standard_bundle=target_standard_bundle,
    )
    standard_row = next((row for row in rows if row["entity_kind"] == "standard"), None)
    return {
        "standard": standard_row["normalized_values"] if standard_row else {},
        "summary": _summarize_rows(rows),
        "rows": rows,
    }


def _is_standard_template_workbook(workbook_bytes: bytes) -> bool:
    try:
        workbook = load_workbook(io.BytesIO(workbook_bytes), read_only=True)
    except Exception:
        return False
    return all(sheet_name in workbook.sheetnames for sheet_name in REQUIRED_SHEETS)


def _extract_xlsx_table_packages(filename: str, workbook_bytes: bytes) -> tuple[list[dict], list[dict]]:
    try:
        workbook = load_workbook(io.BytesIO(workbook_bytes), data_only=True)
    except (BadZipFile, InvalidFileException, OSError) as error:
        raise ValueError("Uploaded file is not a valid .xlsx workbook") from error

    packages: list[dict] = []
    chunks: list[dict] = []
    for sheet in workbook.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        header_index = _find_header_row_index(rows)
        if header_index is None:
            continue
        headers = [_trim_text(value) for value in rows[header_index]]
        table_rows = []
        for source_row_number, row_values in enumerate(rows[header_index + 1 :], start=header_index + 2):
            values = {
                headers[index]: _cell_value(row_values[index] if index < len(row_values) else None)
                for index in range(len(headers))
                if headers[index]
            }
            if _row_is_empty(values):
                continue
            table_rows.append({"source_row_number": source_row_number, "values": values})
        if not table_rows:
            continue
        package = {
            "filename": filename,
            "source_kind": "table",
            "sheet_name": sheet.title,
            "page_no": None,
            "table_index": 1,
            "title": sheet.title,
            "heading_path": [sheet.title],
            "headers": headers,
            "rows": table_rows,
            "context_before": "",
            "context_after": "",
        }
        packages.append(package)
        chunks.append(_table_package_to_chunk(len(chunks) + 1, package))
    return packages, chunks


def _extract_docx_table_packages(filename: str, document_bytes: bytes) -> tuple[list[dict], list[dict]]:
    try:
        from docx import Document
    except ImportError as error:  # pragma: no cover - deployment dependency guard
        raise ValueError("python-docx is required to import .docx standards") from error

    try:
        document = Document(io.BytesIO(document_bytes))
    except Exception as error:
        raise ValueError("Uploaded file is not a valid .docx document") from error

    packages: list[dict] = []
    chunks: list[dict] = []
    heading_path: list[str] = []
    recent_paragraphs: list[str] = []
    table_index = 0
    content_iter = document.iter_inner_content() if hasattr(document, "iter_inner_content") else [*document.paragraphs, *document.tables]
    for block in content_iter:
        if hasattr(block, "rows"):
            table_index += 1
            raw_rows = [[_normalize_cell_text(cell.text) for cell in row.cells] for row in block.rows]
            header_index = _find_header_row_index(raw_rows)
            if header_index is None:
                continue
            headers = raw_rows[header_index]
            table_rows = []
            for offset, row_values in enumerate(raw_rows[header_index + 1 :], start=header_index + 2):
                values = {
                    headers[index]: _cell_value(row_values[index] if index < len(row_values) else None)
                    for index in range(len(headers))
                    if headers[index]
                }
                if _row_is_empty(values):
                    continue
                table_rows.append({"source_row_number": offset, "values": values})
            if not table_rows:
                continue
            title = heading_path[-1] if heading_path else (recent_paragraphs[-1] if recent_paragraphs else f"表格 {table_index}")
            package = {
                "filename": filename,
                "source_kind": "table",
                "sheet_name": None,
                "page_no": None,
                "table_index": table_index,
                "title": title,
                "heading_path": heading_path[:],
                "headers": headers,
                "rows": table_rows,
                "context_before": "\n".join(recent_paragraphs[-3:]),
                "context_after": "",
            }
            packages.append(package)
            chunks.append(_table_package_to_chunk(len(chunks) + 1, package))
            continue

        text = _normalize_cell_text(getattr(block, "text", ""))
        if not text:
            continue
        style_name = getattr(getattr(block, "style", None), "name", "") or ""
        heading_level = _heading_level(style_name)
        if heading_level:
            heading_path = [*heading_path[: heading_level - 1], text]
        recent_paragraphs.append(text)
        if len(recent_paragraphs) > 8:
            recent_paragraphs = recent_paragraphs[-8:]
        chunks.append(
            {
                "chunk_no": len(chunks) + 1,
                "source_kind": "text",
                "sheet_name": None,
                "page_no": None,
                "table_index": None,
                "heading_path": heading_path[:],
                "content": text,
                "metadata": {},
            }
        )
    return packages, chunks


def _extract_pdf_context_packages(filename: str, pdf_bytes: bytes) -> tuple[list[dict], list[dict]]:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        return [], [
            {
                "chunk_no": 1,
                "source_kind": "text",
                "sheet_name": None,
                "page_no": 1,
                "table_index": None,
                "heading_path": [],
                "content": "PDF text extraction is not available in this deployment. Upload .docx or .xlsx for table-accurate import.",
                "metadata": {"parser_warning": "pypdf_not_installed"},
            }
        ]
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
    except Exception as error:
        raise ValueError("Uploaded file is not a valid PDF document") from error
    chunks = []
    for page_index, page in enumerate(reader.pages[:20], start=1):
        text = _normalize_cell_text(page.extract_text() or "")
        if not text:
            continue
        chunks.append(
            {
                "chunk_no": len(chunks) + 1,
                "source_kind": "text",
                "sheet_name": None,
                "page_no": page_index,
                "table_index": None,
                "heading_path": [],
                "content": text[:4000],
                "metadata": {"pdf_text_only": True},
            }
        )
    return [], chunks


def _build_ai_candidate_rows(filename: str, packages: list[dict], *, target_mode: str) -> list[dict]:
    settings = get_ai_settings_secret()
    if not _llm_enabled(settings) or not packages:
        return []

    prompt = {
        "task": "从工程标准 Word/Excel 表格中抽取标准库草稿。只返回 JSON。",
        "target_mode": target_mode,
        "source_filename": filename,
        "allowed_entity_kinds": [
            "standard",
            "pbs_level",
            "tag_class",
            "tag_attribute",
            "equipment_class",
            "equipment_attribute",
            "tag_equipment_class_relationship",
            "document_type",
            "document_attribute",
        ],
        "output_schema": {
            "items": [
                {
                    "entity_kind": "one allowed_entity_kinds value",
                    "confidence": "0..1 number",
                    "values": "object matching the nearest import fields",
                    "evidence": "array of filename/sheet_name/page_no/table_index/row_number/column_name/source_text",
                }
            ]
        },
        "tables": _compact_table_packages_for_ai(packages),
    }
    try:
        text = complete_chat_text(
            settings,
            system_prompt="You are a standards-library table extraction engine. Output valid JSON only.",
            user_prompt=json.dumps(prompt, ensure_ascii=False),
        )
        payload = json.loads(text)
    except Exception:
        return []
    items = payload.get("items") if isinstance(payload, dict) else None
    if not isinstance(items, list):
        return []
    rows = []
    for index, item in enumerate(items, start=1):
        if not isinstance(item, dict) or item.get("entity_kind") not in {
            "standard",
            "pbs_level",
            "tag_class",
            "tag_attribute",
            "equipment_class",
            "equipment_attribute",
            "tag_equipment_class_relationship",
            "document_type",
            "document_attribute",
            "discipline",
            "discipline_document_type",
            "class_document_requirement",
        }:
            continue
        values = item.get("values") if isinstance(item.get("values"), dict) else {}
        rows.append(
            {
                "id": None,
                "row_number": index,
                "source_kind": "table",
                "sheet_name": None,
                "page_no": None,
                "table_index": None,
                "source_row_number": index,
                "entity_kind": item["entity_kind"],
                "values": values,
                "confidence": max(0, min(float(item.get("confidence") or 0.5), 1)),
                "evidence": item.get("evidence") if isinstance(item.get("evidence"), list) else [],
            }
        )
    return rows


def _packages_to_import_rows(filename: str, packages: list[dict]) -> list[dict]:
    rows: list[dict] = []
    source_context = _build_source_table_context(packages)
    for package in packages:
        table_kind = _infer_table_entity_kind(package)
        if table_kind is None:
            continue
        for table_row in package["rows"]:
            values = _map_table_row_values(table_row["values"], table_kind, package, source_context)
            if _row_is_empty(values):
                continue
            rows.append(
                {
                    "id": None,
                    "row_number": len(rows) + 1,
                    "source_kind": "table",
                    "sheet_name": package.get("sheet_name"),
                    "page_no": package.get("page_no"),
                    "table_index": package.get("table_index"),
                    "source_row_number": table_row["source_row_number"],
                    "entity_kind": table_kind,
                    "values": values,
                    "confidence": 0.72,
                    "evidence": _table_row_evidence(filename, package, table_row),
                }
            )
    return rows


def _fallback_standard_row(
    filename: str,
    file_ext: str,
    target_mode: str,
    target_standard_bundle: dict | None,
) -> dict:
    if target_mode == "merge" and target_standard_bundle:
        standard = target_standard_bundle["standard"]
        values = {
            "code": standard["code"],
            "name": standard["name"],
            "version_label": standard.get("version_label"),
            "status": standard.get("status", "active"),
            "thumbnail_url": standard.get("thumbnail_url"),
            "metadata_json": _json_text(standard.get("metadata") or {}),
        }
    elif file_ext == ".pdf":
        values = {"code": "", "name": "", "version_label": "", "status": "draft", "thumbnail_url": "", "metadata_json": "{}"}
    else:
        stem = os.path.splitext(filename)[0]
        values = {"code": _normalize_code(stem)[:64], "name": stem, "version_label": "", "status": "draft", "thumbnail_url": "", "metadata_json": "{}"}
    return {
        "id": None,
        "row_number": 1,
        "source_kind": "text" if file_ext == ".pdf" else "manual",
        "sheet_name": None,
        "page_no": 1 if file_ext == ".pdf" else None,
        "table_index": None,
        "source_row_number": 1,
        "entity_kind": "standard",
        "values": values,
        "confidence": 0.35 if file_ext == ".pdf" else 0.55,
        "evidence": [
            {
                "filename": filename,
                "sheet_name": None,
                "page_no": 1 if file_ext == ".pdf" else None,
                "table_index": None,
                "row_number": 1,
                "column_name": None,
                "source_text": "系统根据源文件生成标准草稿占位；请在预览中确认编码和名称。",
            }
        ],
    }


def get_standard_import_job_detail(
    job_id: str,
    *,
    status: str | None = None,
    item_type: str | None = None,
    source_table: str | None = None,
    page: int = 1,
    page_size: int = DEFAULT_PAGE_SIZE,
) -> dict:
    normalized_page_size = max(1, min(page_size, MAX_PAGE_SIZE))
    normalized_page = max(1, page)
    filter_status = status if status in ROW_STATUS_VALUES else None

    job = fetch_one(
        """
        SELECT id, filename, file_ext, file_size, checksum_sha256, target_mode,
               summary, status, target_standard_id, source_standard_code,
               committed_at, created_at, updated_at
        FROM standard_import_job
        WHERE id = %s
        """,
        (job_id,),
    )
    if job is None:
        raise ValueError("Standard import job not found")

    filter_sql = ""
    filter_params: list[Any] = []
    if filter_status:
        filter_sql = " AND r.status = %s"
        filter_params.append(filter_status)
    if item_type:
        filter_sql += " AND r.entity_kind = %s"
        filter_params.append(item_type)
    if source_table:
        filter_sql += " AND (r.sheet_name = %s OR r.table_index::text = %s)"
        filter_params.extend([source_table, source_table])

    total_rows_result = fetch_one(
        f"""
        SELECT COUNT(*)::int AS total_rows
        FROM standard_import_row r
        WHERE r.job_id = %s
        {filter_sql}
        """,
        (job_id, *filter_params),
    )
    total_rows = total_rows_result["total_rows"] if total_rows_result else 0
    total_pages = max(1, math.ceil(total_rows / normalized_page_size)) if total_rows else 1
    offset = (normalized_page - 1) * normalized_page_size
    rows = fetch_all(
        f"""
        SELECT id, row_number, source_kind, sheet_name, page_no, table_index,
               source_row_number, entity_kind, values, normalized_values, issues,
               status, action, confidence, evidence, created_at, updated_at
        FROM standard_import_row r
        WHERE r.job_id = %s
        {filter_sql}
        ORDER BY r.row_number
        LIMIT %s OFFSET %s
        """,
        (job_id, *filter_params, normalized_page_size, offset),
    )

    return {
        "job_id": str(job["id"]),
        "filename": job["filename"],
        "file_ext": job["file_ext"],
        "file_size": job["file_size"],
        "checksum_sha256": job["checksum_sha256"],
        "target_mode": job["target_mode"],
        "summary": job["summary"],
        "items": [_serialize_import_row(row) for row in rows],
        "rows": [_serialize_import_row(row) for row in rows],
        "page": normalized_page,
        "page_size": normalized_page_size,
        "total_pages": total_pages,
        "status": job["status"],
        "target_standard_id": str(job["target_standard_id"]) if job.get("target_standard_id") else None,
        "source_standard_code": job.get("source_standard_code"),
        "created_at": job["created_at"],
        "updated_at": job["updated_at"],
        "committed_at": job["committed_at"],
    }


def patch_standard_import_job(job_id: str, payload: dict) -> dict:
    action = _trim_text(payload.get("conflict_action"))
    if action not in CONFLICT_ACTIONS:
        raise ValueError("conflict_action must be one of create_copy, merge_update, or skip")
    item_action = "update" if action == "merge_update" else "skip"
    if action == "create_copy":
        item_action = "create"

    rows = _load_standard_import_rows(job_id)
    action_by_row = {row["id"]: item_action for row in rows if row["status"] == "conflict"}
    if action == "create_copy":
        code_override = _normalize_code(payload.get("code_override"))
        if not code_override:
            raise ValueError("code_override is required when conflict_action is create_copy")
        if _standard_code_exists(code_override):
            raise ValueError("code_override already exists")
        for row in rows:
            if row["entity_kind"] == "standard":
                row["values"] = {**row["values"], "code": code_override}

    _replace_standard_import_rows(job_id, rows, action_by_row=action_by_row)
    return get_standard_import_job_detail(job_id)


def patch_standard_import_item(job_id: str, item_id: str, payload: dict) -> dict:
    _ensure_standard_import_job_exists(job_id)
    rows = _load_standard_import_rows(job_id)
    if not any(row["id"] == item_id for row in rows):
        raise ValueError("Standard import item not found")
    action_by_row = {
        row["id"]: row.get("action")
        for row in rows
        if row.get("action") in ITEM_ACTIONS
    }
    for row in rows:
        if row["id"] != item_id:
            continue
        if "values" in payload and isinstance(payload["values"], dict):
            row["values"] = {**row["values"], **payload["values"]}
        if "action" in payload:
            action = payload.get("action")
            if action is not None and action not in ITEM_ACTIONS:
                raise ValueError("action must be one of create, update, or skip")
            if action is None:
                action_by_row.pop(row["id"], None)
            else:
                action_by_row[row["id"]] = action

    _replace_standard_import_rows(job_id, rows, action_by_row=action_by_row)
    refreshed = get_standard_import_job_detail(job_id)
    refreshed["item"] = next((item for item in refreshed["items"] if item["id"] == item_id), None)
    return refreshed


def commit_standard_import_job(job_id: str) -> dict:
    job = _ensure_standard_import_job_exists(job_id)
    if job["status"] == "committed":
        raise ValueError("Standard import job has already been committed")

    rows = _load_standard_import_rows(job_id)
    if not rows:
        raise ValueError("Standard import job has no rows")
    if any(row["status"] == "error" for row in rows):
        raise ValueError("Cannot commit while error rows still exist")
    if any(row["status"] == "conflict" for row in rows):
        raise ValueError("Conflict rows must be resolved before commit")

    standard_row = next((row for row in rows if row["entity_kind"] == "standard"), None)
    if standard_row is None and job["target_mode"] == "new":
        raise ValueError("Standard import job is missing the standard row")
    standard = standard_row["normalized_values"] if standard_row else {}

    created_count = 0
    updated_count = 0
    skipped_count = 0
    target_standard_id: str | None = None

    with get_connection() as connection:
        with connection.cursor() as cursor:
            if job["target_mode"] == "merge":
                target_standard_id = str(job["target_standard_id"])
                if standard_row and standard_row.get("action") == "update":
                    _update_standard(cursor, target_standard_id, standard)
                    updated_count += 1
                elif standard_row and standard_row.get("action") == "skip":
                    skipped_count += 1
            else:
                if standard_row and standard_row.get("action") == "skip":
                    skipped_count += 1
                    _mark_import_job_committed(cursor, job_id, None, job.get("summary"))
                    connection.commit()
                    return _commit_result(job_id, None, created_count, updated_count, skipped_count)
                existing_standard = _fetch_standard_by_code(cursor, standard["code"])
                if existing_standard is not None:
                    if standard_row and standard_row.get("action") == "update":
                        target_standard_id = str(existing_standard["id"])
                        _update_standard(cursor, target_standard_id, standard)
                        updated_count += 1
                    else:
                        raise ValueError("Standard code already exists; update the code or choose update")
                else:
                    target_standard_id = _insert_standard(cursor, standard, standard["code"])
                    created_count += 1

            tag_class_ids = _upsert_classes(
                cursor,
                target_standard_id,
                _committable_values(rows, "tag_class"),
                applies_to="tag",
            )
            equipment_class_refs = _upsert_classes(
                cursor,
                target_standard_id,
                _committable_values(rows, "equipment_class"),
                applies_to="equipment",
            )
            document_type_rows = _committable_values(rows, "document_type")
            document_type_ids = _upsert_classes(
                cursor,
                target_standard_id,
                document_type_rows,
                applies_to="document",
            )
            discipline_refs = _upsert_disciplines(
                cursor,
                target_standard_id,
                _committable_values(rows, "discipline"),
            )
            pbs_counts = _upsert_pbs_levels(
                cursor,
                target_standard_id,
                _committable_values(rows, "pbs_level"),
            )
            tag_attribute_counts = _upsert_attributes(
                cursor,
                target_standard_id,
                _committable_values(rows, "tag_attribute"),
                applies_to="tag",
                owner_ids=tag_class_ids,
            )
            equipment_attribute_counts = _upsert_attributes(
                cursor,
                target_standard_id,
                _committable_values(rows, "equipment_attribute"),
                applies_to="equipment",
                owner_ids=equipment_class_refs,
            )
            relationship_counts = _upsert_tag_equipment_class_relationships(
                cursor,
                target_standard_id,
                _committable_values(rows, "tag_equipment_class_relationship"),
                tag_class_ids=tag_class_ids,
                equipment_class_refs=equipment_class_refs,
            )
            document_attribute_counts = _upsert_attributes(
                cursor,
                target_standard_id,
                _committable_values(rows, "document_attribute"),
                applies_to="document",
                owner_ids=document_type_ids,
            )
            document_type_cfihos_refs = _document_type_cfihos_refs(document_type_ids, document_type_rows)
            discipline_document_type_counts = _upsert_discipline_document_types(
                cursor,
                target_standard_id,
                _committable_values(rows, "discipline_document_type"),
                discipline_refs=discipline_refs,
                document_type_ids=document_type_ids,
                document_type_cfihos_refs=document_type_cfihos_refs,
            )
            class_document_requirement_counts = _upsert_class_document_requirements(
                cursor,
                target_standard_id,
                _committable_values(rows, "class_document_requirement"),
                tag_class_ids=tag_class_ids,
                equipment_class_refs=equipment_class_refs,
                document_type_ids=document_type_ids,
                document_type_cfihos_refs=document_type_cfihos_refs,
            )
            skipped_count += sum(1 for row in rows if row.get("action") == "skip" and row["entity_kind"] != "standard")
            created_count += (
                pbs_counts[0]
                + tag_attribute_counts[0]
                + equipment_attribute_counts[0]
                + relationship_counts[0]
                + document_attribute_counts[0]
                + discipline_document_type_counts[0]
                + class_document_requirement_counts[0]
            )
            updated_count += (
                pbs_counts[1]
                + tag_attribute_counts[1]
                + equipment_attribute_counts[1]
                + relationship_counts[1]
                + document_attribute_counts[1]
                + discipline_document_type_counts[1]
                + class_document_requirement_counts[1]
            )
            created_count += sum(1 for item in tag_class_ids.values() if item["created"])
            updated_count += sum(1 for item in tag_class_ids.values() if not item["created"])
            created_count += sum(1 for item in equipment_class_refs.values() if item["created"])
            updated_count += sum(1 for item in equipment_class_refs.values() if not item["created"])
            created_count += sum(1 for item in document_type_ids.values() if item["created"])
            updated_count += sum(1 for item in document_type_ids.values() if not item["created"])
            created_count += discipline_refs["created_count"]
            updated_count += discipline_refs["updated_count"]

            _mark_import_job_committed(cursor, job_id, target_standard_id, job.get("summary"))
        connection.commit()

    return _commit_result(job_id, target_standard_id, created_count, updated_count, skipped_count)


def _parse_standard_import_workbook(workbook_bytes: bytes) -> list[dict]:
    workbook = load_workbook(io.BytesIO(workbook_bytes), data_only=True)
    missing_sheets = [sheet_name for sheet_name in REQUIRED_SHEETS if sheet_name not in workbook.sheetnames]
    if missing_sheets:
        raise ValueError("Workbook must contain sheets: " + ", ".join(missing_sheets))

    row_values: list[dict] = []
    import_row_number = 1
    standard_rows = _read_sheet_rows(workbook, STANDARD_SHEET, STANDARD_HEADERS, "standard")
    if len(standard_rows) != 1:
        raise ValueError("Workbook must contain exactly one row in the standard sheet")

    sheet_specs = [
        (STANDARD_SHEET, STANDARD_HEADERS, "standard"),
        (PBS_LEVEL_SHEET, PBS_LEVEL_HEADERS, "pbs_level"),
        (TAG_CLASS_SHEET, CLASS_HEADERS, "tag_class"),
        (TAG_ATTRIBUTE_SHEET, TAG_ATTRIBUTE_HEADERS, "tag_attribute"),
        (DOCUMENT_TYPE_SHEET, DOCUMENT_TYPE_HEADERS, "document_type"),
        (DOCUMENT_ATTRIBUTE_SHEET, DOCUMENT_ATTRIBUTE_HEADERS, "document_attribute"),
    ]
    for optional_spec in [
        (EQUIPMENT_CLASS_SHEET, CLASS_HEADERS, "equipment_class"),
        (EQUIPMENT_ATTRIBUTE_SHEET, TAG_ATTRIBUTE_HEADERS, "equipment_attribute"),
        (TAG_EQUIPMENT_CLASS_RELATIONSHIP_SHEET, TAG_EQUIPMENT_CLASS_RELATIONSHIP_HEADERS, "tag_equipment_class_relationship"),
    ]:
        if optional_spec[0] in workbook.sheetnames:
            sheet_specs.append(optional_spec)

    for sheet_name, headers, entity_kind in sheet_specs:
        for row in _read_sheet_rows(workbook, sheet_name, headers, entity_kind):
            row_values.append({**row, "row_number": import_row_number})
            import_row_number += 1
    return row_values


def _find_header_row_index(rows: list[tuple | list]) -> int | None:
    for index, row in enumerate(rows[:20]):
        normalized = [_trim_text(value) for value in row]
        non_empty = [value for value in normalized if value]
        if len(non_empty) < 2:
            continue
        recognized = sum(1 for value in non_empty if _canonical_header(value) is not None)
        if recognized >= 2 or any(token in "".join(non_empty).lower() for token in ["code", "编码", "名称", "属性"]):
            return index
    return None


def _canonical_header(header: Any) -> str | None:
    text = _normalize_header(header)
    if not text:
        return None
    compact = re.sub(r"[\s_\-:：/（）()]+", "", text).lower()
    for canonical, aliases in HEADER_ALIASES.items():
        for alias in aliases:
            alias_compact = re.sub(r"[\s_\-:：/（）()]+", "", alias).lower()
            if compact == alias_compact or alias_compact in compact:
                return canonical
    return None


def _infer_table_entity_kind(package: dict) -> str | None:
    if _is_cfihos_property_dictionary(package):
        return None

    source_table = _source_table_name(package)
    if source_table in UNSUPPORTED_SOURCE_TABLES:
        return None
    if source_table in KNOWN_SOURCE_TABLE_KINDS:
        return KNOWN_SOURCE_TABLE_KINDS[source_table]

    title = " ".join([*(package.get("heading_path") or []), str(package.get("title") or "")]).lower()
    headers = [_canonical_header(header) for header in package.get("headers", [])]
    header_set = {header for header in headers if header}

    if "standard_code" in header_set or ("标准" in title and {"code", "name"}.issubset(header_set)):
        return "standard"
    if "level_no" in header_set and ("pbs" in title or "层级" in title):
        return "pbs_level"
    if {"value_type", "is_required"} & header_set or "属性" in title:
        if "document" in title or "文档" in title or "图纸" in title or "owner_document_type_code" in header_set:
            return "document_attribute"
        return "tag_attribute"
    if "allowed_extensions" in header_set or "文档" in title or "图纸" in title:
        return "document_type"
    if {"code", "name"}.issubset(header_set) and ("位号" in title or "class" in title or "类别" in title or "类型" in title):
        return "tag_class"
    return None


def _map_table_row_values(row_values: dict, entity_kind: str, package: dict, source_context: dict | None = None) -> dict:
    known_values = _map_known_source_table_row_values(row_values, entity_kind, package, source_context or {})
    if known_values is not None:
        return known_values

    mapped: dict[str, Any] = {}
    for original_key, value in row_values.items():
        canonical = _canonical_header(original_key)
        if canonical:
            mapped[canonical] = value

    if entity_kind == "standard":
        return {
            "code": mapped.get("standard_code") or mapped.get("code"),
            "name": mapped.get("name"),
            "version_label": mapped.get("version_label"),
            "status": mapped.get("status") or "draft",
            "thumbnail_url": "",
            "metadata_json": "{}",
        }
    if entity_kind == "pbs_level":
        return {
            "level_no": mapped.get("level_no"),
            "code": mapped.get("code"),
            "name": mapped.get("name"),
            "description": mapped.get("description"),
        }
    if entity_kind in {"tag_class", "equipment_class"}:
        return {
            "code": mapped.get("code"),
            "name": mapped.get("name"),
            "parent_code": mapped.get("parent_code"),
            "level_no": mapped.get("level_no") or 1,
            "description": mapped.get("description"),
            "status": mapped.get("status") or "active",
        }
    if entity_kind == "tag_equipment_class_relationship":
        return {
            "tag_class_code": mapped.get("tag_class_code"),
            "equipment_class_code": mapped.get("equipment_class_code"),
            "reason": mapped.get("description"),
            "status": mapped.get("status") or "active",
        }
    if entity_kind == "document_type":
        return {
            "code": mapped.get("code"),
            "name": mapped.get("name"),
            "parent_code": mapped.get("parent_code"),
            "level_no": mapped.get("level_no") or 1,
            "description": mapped.get("description"),
            "status": mapped.get("status") or "active",
            "allowed_extensions": mapped.get("allowed_extensions"),
            "metadata_json": "{}",
        }

    owner_field = "owner_document_type_code" if entity_kind == "document_attribute" else "owner_class_code"
    return {
        owner_field: mapped.get(owner_field),
        "group_name": mapped.get("group_name"),
        "code": mapped.get("code"),
        "name": mapped.get("name"),
        "value_type": _normalize_value_type(mapped.get("value_type")),
        "is_required": mapped.get("is_required"),
        "unit_family": mapped.get("unit_family"),
        "enum_options": mapped.get("enum_options"),
        "description": mapped.get("description"),
        "sort_order": mapped.get("sort_order") or 0,
        "status": mapped.get("status") or "active",
    }


def _source_table_name(package: dict) -> str:
    return _trim_text(package.get("sheet_name") or package.get("title")).lower()


def _is_cfihos_property_dictionary(package: dict) -> bool:
    if _source_table_name(package) != "property":
        return False
    headers = {_normalize_header(header).lower() for header in package.get("headers", [])}
    return {"cfihos unique code", "property name", "property data type"}.issubset(headers)


def _build_source_table_context(packages: list[dict]) -> dict:
    discipline_by_cfihos: dict[str, str] = {}
    discipline_by_code: dict[str, str] = {}
    document_type_by_cfihos: dict[str, str] = {}
    equipment_class_by_name: dict[str, str] = {}
    tag_class_by_name: dict[str, str] = {}
    property_by_code: dict[str, dict] = {}
    picklist_options_by_name: dict[str, list[str]] = {}
    for package in packages:
        source_table = _source_table_name(package)
        for table_row in package.get("rows", []):
            values = table_row.get("values") or {}
            if source_table == "discipline":
                cfihos_code = _normalize_code(_source_value(values, "CFIHOS unique code", "discipline CFIHOS unique code"))
                discipline_code = _normalize_code(_source_value(values, "discipline code"))
                if cfihos_code and discipline_code:
                    discipline_by_cfihos[_code_key(cfihos_code)] = discipline_code
                if discipline_code:
                    discipline_by_code[_code_key(discipline_code)] = discipline_code
            elif source_table == "document type":
                cfihos_code = _normalize_code(_source_value(values, "CFIHOS unique code", "document type CFIHOS unique code"))
                document_type_code = _normalize_code(_source_value(values, "document type short code", "CFIHOS unique code"))
                if cfihos_code and document_type_code:
                    document_type_by_cfihos[_code_key(cfihos_code)] = document_type_code
            elif source_table == "equipment class":
                code = _normalize_code(_source_value(values, "equipment class CFIHOS unique code"))
                name = _trim_text(_source_value(values, "equipment class name")).lower()
                if code and name:
                    equipment_class_by_name[name] = code
            elif source_table == "tag class":
                code = _normalize_code(_source_value(values, "CFIHOS unique code", "tag class CFIHOS unique code"))
                name = _trim_text(_source_value(values, "tag class name")).lower()
                if code and name:
                    tag_class_by_name[name] = code
            elif source_table == "property":
                code = _normalize_code(_source_value(values, "CFIHOS unique code"))
                if code:
                    property_by_code[_code_key(code)] = {
                        "code": code,
                        "name": _source_value(values, "property name"),
                        "value_type": _normalize_cfihos_property_value_type(values),
                        "unit_family": _source_value(values, "unit of measure dimension code", "unit of measure dimension code CFIHOS unique code"),
                        "description": _source_value(values, "property definition"),
                        "picklist_name": _source_value(values, "property picklist name"),
                    }
            elif source_table == "property picklist values":
                picklist_name = _trim_text(_source_value(values, "property picklist name")).lower()
                option = _source_value(values, "property picklist value code", "property picklist value description")
                if picklist_name and _blank_to_none(option) is not None:
                    picklist_options_by_name.setdefault(picklist_name, [])
                    option_text = _trim_text(option)
                    if option_text not in picklist_options_by_name[picklist_name]:
                        picklist_options_by_name[picklist_name].append(option_text)
    return {
        "discipline_by_cfihos": discipline_by_cfihos,
        "discipline_by_code": discipline_by_code,
        "document_type_by_cfihos": document_type_by_cfihos,
        "equipment_class_by_name": equipment_class_by_name,
        "tag_class_by_name": tag_class_by_name,
        "property_by_code": property_by_code,
        "picklist_options_by_name": picklist_options_by_name,
    }


def _source_value(row_values: dict, *headers: str) -> Any:
    by_key = {_normalize_header(key).lower(): value for key, value in row_values.items()}
    for header in headers:
        value = by_key.get(_normalize_header(header).lower())
        if _blank_to_none(value) is not None:
            return value
    return None


def _normalize_context_code(value: Any, *, default: str | None = None) -> str | None:
    text = _trim_text(value).lower()
    if not text:
        return default
    normalized = re.sub(r"[^a-z0-9]+", "_", text).strip("_")
    return normalized or default


def _normalize_lifecycle_phase(value: Any) -> str:
    text = _trim_text(value).lower()
    aliases = {
        "": "unspecified",
        "during project": "project",
        "during projects": "project",
        "at handover": "handover",
        "handover": "handover",
        "not applicable": "not_applicable",
        "n/a": "not_applicable",
        "na": "not_applicable",
    }
    return aliases.get(text, _normalize_context_code(text, default="unspecified") or "unspecified")


def _map_known_source_table_row_values(row_values: dict, entity_kind: str, package: dict, source_context: dict) -> dict | None:
    source_table = _source_table_name(package)
    if source_table not in KNOWN_SOURCE_TABLE_KINDS:
        return None

    if source_table == "document type" and entity_kind == "document_type":
        return {
            "code": _source_value(row_values, "document type short code", "CFIHOS unique code"),
            "name": _source_value(row_values, "document type name"),
            "parent_code": "",
            "level_no": 1,
            "description": _source_value(row_values, "document type description"),
            "status": "active",
            "allowed_extensions": "",
            "metadata_json": _json_text(
                {
                    "cfihos_unique_code": _source_value(row_values, "CFIHOS unique code"),
                    "classification": _source_value(row_values, "document type classification"),
                    "synonym_name": _source_value(row_values, "document type synonym name"),
                }
            ),
        }

    if source_table == "discipline" and entity_kind == "discipline":
        return {
            "cfihos_unique_code": _source_value(row_values, "CFIHOS unique code", "discipline CFIHOS unique code"),
            "code": _source_value(row_values, "discipline code"),
            "name": _source_value(row_values, "discipline name"),
            "description": _source_value(row_values, "discipline description"),
            "status": "active",
            "metadata_json": _json_text(
                {
                    "source_table": source_table,
                }
            ),
        }

    if source_table == "discipline document type" and entity_kind == "discipline_document_type":
        discipline_cfihos_code = _normalize_code(_source_value(row_values, "discipline CFIHOS unique code"))
        document_type_cfihos_code = _normalize_code(_source_value(row_values, "document type CFIHOS unique code"))
        raw_asset_scope = _source_value(row_values, "asset type reference")
        raw_delivery_timing = _source_value(row_values, "native file delivery timing")
        return {
            "cfihos_unique_code": _source_value(row_values, "discipline document type CFIHOS unique code"),
            "discipline_cfihos_unique_code": discipline_cfihos_code,
            "discipline_code": _source_value(row_values, "discipline code")
            or (source_context.get("discipline_by_cfihos") or {}).get(_code_key(discipline_cfihos_code)),
            "document_type_cfihos_unique_code": document_type_cfihos_code,
            "document_type_code": _source_value(row_values, "document type short code")
            or (source_context.get("document_type_by_cfihos") or {}).get(_code_key(document_type_cfihos_code)),
            "document_type_name": _source_value(row_values, "document type name"),
            "short_code": _source_value(row_values, "discipline document type short code"),
            "asset_scope": _normalize_context_code(raw_asset_scope),
            "representation_type": _blank_to_none(_source_value(row_values, "representation type")),
            "native_file_delivery_timing": _blank_to_none(raw_delivery_timing),
            "perspective": "standard",
            "lifecycle_phase": _normalize_lifecycle_phase(raw_delivery_timing),
            "status": "active",
            "metadata_json": _json_text(
                {
                    "raw_asset_scope": raw_asset_scope,
                    "raw_native_file_delivery_timing": raw_delivery_timing,
                    "discipline_name": _source_value(row_values, "discipline name"),
                    "document_type_description": _source_value(row_values, "document type description"),
                    "source_table": source_table,
                }
            ),
        }

    if source_table == "equipment class" and entity_kind == "equipment_class":
        parent_name = _trim_text(_source_value(row_values, "parent equipment class name")).lower()
        return {
            "code": _source_value(row_values, "equipment class CFIHOS unique code"),
            "name": _source_value(row_values, "equipment class name"),
            "parent_code": (source_context.get("equipment_class_by_name") or {}).get(parent_name, ""),
            "level_no": "",
            "description": _source_value(row_values, "equipment class definition"),
            "status": "active",
        }

    if source_table == "tag class" and entity_kind == "tag_class":
        parent_name = _trim_text(_source_value(row_values, "parent tag class name")).lower()
        return {
            "code": _source_value(row_values, "CFIHOS unique code", "tag class CFIHOS unique code"),
            "name": _source_value(row_values, "tag class name"),
            "parent_code": (source_context.get("tag_class_by_name") or {}).get(parent_name, ""),
            "level_no": "",
            "description": _source_value(row_values, "tag class definition"),
            "status": "active",
        }

    if source_table == "data dictionary" and entity_kind == "equipment_attribute":
        return _map_cfihos_equipment_dictionary_row(row_values)

    if source_table == "equipment class property" and entity_kind == "equipment_attribute":
        return _map_cfihos_class_property_row(
            row_values,
            owner_header="equipment class CFIHOS unique code",
            required_header="property relevant for equipment indicator",
            source_context=source_context,
        )

    if source_table == "tag class property" and entity_kind == "tag_attribute":
        return _map_cfihos_class_property_row(
            row_values,
            owner_header="tag class CFIHOS unique code",
            required_header=None,
            source_context=source_context,
        )

    if source_table == "tag equipment class relationshi" and entity_kind == "tag_equipment_class_relationship":
        return {
            "tag_class_code": _source_value(row_values, "tag class CFIHOS unique code"),
            "tag_class_name": _source_value(row_values, "tag class name"),
            "equipment_class_code": _source_value(row_values, "equipment class CFIHOS unique code"),
            "equipment_class_name": _source_value(row_values, "equipment class name"),
            "reason": _source_value(row_values, "tag or equipment class relationship reason for mapping"),
            "status": "active",
        }

    if source_table == "document required per class" and entity_kind == "class_document_requirement":
        document_type_cfihos_code = _normalize_code(_source_value(row_values, "document type CFIHOS unique code"))
        raw_asset_scope = _source_value(row_values, "asset type reference")
        return {
            "cfihos_unique_code": _source_value(row_values, "source standard document and data requirement CFIHOS unique code"),
            "class_code": _source_value(row_values, "tag or equipment class CFIHOS unique code"),
            "class_name": _source_value(row_values, "tag or equipment class name"),
            "asset_scope": _normalize_context_code(raw_asset_scope),
            "source_standard_cfihos_code": _source_value(row_values, "source standard CFIHOS unique code"),
            "source_standard_code": _source_value(row_values, "source standard code"),
            "document_type_cfihos_unique_code": document_type_cfihos_code,
            "document_type_code": (source_context.get("document_type_by_cfihos") or {}).get(_code_key(document_type_cfihos_code)),
            "document_type_name": _source_value(row_values, "document type name"),
            "perspective": "standard",
            "lifecycle_phase": "unspecified",
            "status": "active",
            "metadata_json": _json_text(
                {
                    "raw_asset_scope": raw_asset_scope,
                    "class_name": _source_value(row_values, "tag or equipment class name"),
                    "document_type_name": _source_value(row_values, "document type name"),
                    "source_table": source_table,
                }
            ),
        }

    return None


def _map_cfihos_equipment_dictionary_row(row_values: dict) -> dict:
    if _trim_text(_source_value(row_values, "object")).lower() != "attribute:":
        return {}
    if _trim_text(_source_value(row_values, "entity name")).lower() != "equipment":
        return {}

    requirement = _trim_text(_source_value(row_values, "identifier / mandatory / optional")).lower()
    is_required = requirement in {"identifier", "mandatory"}
    return {
        "owner_class_code": "",
        "group_name": "CFIHOS equipment fixed attributes",
        "code": _source_value(row_values, "CFIHOS unique code"),
        "name": _source_value(row_values, "property name", "entity attribute name"),
        "value_type": _normalize_cfihos_dictionary_value_type(_source_value(row_values, "format")),
        "is_required": is_required,
        "unit_family": "",
        "enum_options": [],
        "description": _source_value(row_values, "definition"),
        "sort_order": 0,
        "status": "active",
    }


def _map_cfihos_class_property_row(
    row_values: dict,
    *,
    owner_header: str,
    required_header: str | None,
    source_context: dict,
) -> dict:
    property_code = _normalize_code(_source_value(row_values, "property CFIHOS unique code"))
    property_definition = (source_context.get("property_by_code") or {}).get(_code_key(property_code), {})
    picklist_name = _trim_text(property_definition.get("picklist_name")).lower()
    enum_options = (source_context.get("picklist_options_by_name") or {}).get(picklist_name, [])
    return {
        "owner_class_code": _source_value(row_values, owner_header),
        "group_name": "",
        "code": property_code,
        "name": _source_value(row_values, "property name") or property_definition.get("name"),
        "value_type": "enum" if enum_options else property_definition.get("value_type", "string"),
        "is_required": _source_value(row_values, required_header) if required_header else False,
        "unit_family": _source_value(row_values, "SI unit of measure name", "SI unit of measure CFIHOS unique code") or property_definition.get("unit_family"),
        "enum_options": enum_options,
        "description": property_definition.get("description"),
        "sort_order": 0,
        "status": "active",
    }


def _normalize_cfihos_property_value_type(row_values: dict) -> str:
    if _blank_to_none(_source_value(row_values, "property picklist name")) is not None:
        return "enum"
    return _normalize_value_type(_source_value(row_values, "property data type"))


def _normalize_cfihos_dictionary_value_type(value: Any) -> str:
    text = _trim_text(value).lower()
    if "boolean" in text:
        return "boolean"
    if "date" in text:
        return "date"
    if "integer" in text:
        return "integer"
    if "decimal" in text or text == "num" or text.startswith("num "):
        return "number"
    return "string"


def _normalize_value_type(value: Any) -> str:
    text = _trim_text(value).lower()
    aliases = {
        "文本": "string",
        "字符串": "string",
        "text": "string",
        "alphanumeric": "string",
        "数字": "number",
        "数值": "number",
        "整数": "integer",
        "布尔": "boolean",
        "是否": "boolean",
        "日期": "date",
        "枚举": "enum",
        "列表": "enum",
        "json": "json",
    }
    if text in VALID_VALUE_TYPES:
        return text
    return aliases.get(text, "string")


def _table_row_evidence(filename: str, package: dict, table_row: dict) -> list[dict]:
    evidence = []
    for column_name, value in table_row["values"].items():
        if _blank_to_none(value) is None:
            continue
        evidence.append(
            {
                "filename": filename,
                "sheet_name": package.get("sheet_name"),
                "page_no": package.get("page_no"),
                "table_index": package.get("table_index"),
                "row_number": table_row["source_row_number"],
                "column_name": column_name,
                "source_text": _trim_text(value)[:500],
            }
        )
    return evidence[:12]


def _table_package_to_chunk(chunk_no: int, package: dict) -> dict:
    content_lines = ["\t".join(_trim_text(header) for header in package.get("headers", []))]
    for row in package.get("rows", [])[:MAX_AI_TABLE_ROWS]:
        values = row.get("values", {})
        content_lines.append("\t".join(_trim_text(values.get(header)) for header in package.get("headers", [])))
    return {
        "chunk_no": chunk_no,
        "source_kind": package.get("source_kind", "table"),
        "sheet_name": package.get("sheet_name"),
        "page_no": package.get("page_no"),
        "table_index": package.get("table_index"),
        "heading_path": package.get("heading_path") or [],
        "content": "\n".join(content_lines)[:12000],
        "metadata": {"title": package.get("title"), "row_count": len(package.get("rows", []))},
    }


def _compact_table_packages_for_ai(packages: list[dict]) -> list[dict]:
    compact = []
    for package in packages:
        if _is_cfihos_property_dictionary(package):
            continue
        compact.append(
            {
                "filename": package.get("filename"),
                "sheet_name": package.get("sheet_name"),
                "page_no": package.get("page_no"),
                "table_index": package.get("table_index"),
                "title": package.get("title"),
                "heading_path": package.get("heading_path") or [],
                "headers": package.get("headers") or [],
                "context_before": package.get("context_before") or "",
                "rows": package.get("rows", [])[:MAX_AI_TABLE_ROWS],
            }
        )
        if len(compact) >= AI_TABLE_PACKAGE_LIMIT:
            break
    return compact


def _merge_candidate_rows(rows: list[dict]) -> list[dict]:
    merged: list[dict] = []
    seen: set[tuple[str, str, str]] = set()
    for row in rows:
        values = row.get("values", {})
        entity_kind = row.get("entity_kind", "")
        owner = _code_key(values.get("owner_class_code") or values.get("owner_document_type_code")) or "__root__"
        code = _code_key(values.get("standard_code") or values.get("code"))
        if entity_kind == "discipline":
            code = _code_key(values.get("cfihos_unique_code") or values.get("code"))
        elif entity_kind == "discipline_document_type":
            owner = _code_key(values.get("discipline_cfihos_unique_code") or values.get("discipline_code")) or "__root__"
            code = _code_key(
                values.get("cfihos_unique_code")
                or values.get("short_code")
                or "|".join(
                    _trim_text(values.get(field))
                    for field in ["document_type_cfihos_unique_code", "document_type_code", "asset_scope", "representation_type", "native_file_delivery_timing"]
                )
            )
        elif entity_kind == "class_document_requirement":
            owner = _code_key(values.get("class_code")) or "__root__"
            code = _code_key(
                values.get("cfihos_unique_code")
                or "|".join(
                    _trim_text(values.get(field))
                    for field in ["document_type_cfihos_unique_code", "document_type_code", "asset_scope", "source_standard_code"]
                )
            )
        key = (row.get("entity_kind", ""), owner, code or str(row.get("row_number")))
        if key in seen:
            continue
        seen.add(key)
        merged.append(row)
    return merged


def _renumber_import_rows(rows: list[dict]) -> list[dict]:
    return [{**row, "row_number": index} for index, row in enumerate(rows, start=1)]


def _apply_row_actions(rows: list[dict], action_by_row: dict[str, str | None]) -> list[dict]:
    next_rows = []
    for row in rows:
        action = action_by_row.get(str(row["id"])) or action_by_row.get(row["id"])
        if action in ITEM_ACTIONS:
            row = {**row, "action": action}
            if row["status"] == "conflict":
                row = {**row, "status": "ready"}
        next_rows.append(row)
    return next_rows


def _normalize_cell_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").replace("\u3000", " ")).strip()


def _heading_level(style_name: str) -> int | None:
    match = re.search(r"heading\s*(\d+)|标题\s*(\d+)", style_name, flags=re.IGNORECASE)
    if not match:
        return None
    value = match.group(1) or match.group(2)
    try:
        return max(1, min(int(value), 6))
    except ValueError:
        return None


def _llm_enabled(settings: dict) -> bool:
    return bool(
        settings.get("is_enabled")
        and settings.get("base_url")
        and settings.get("model")
        and settings.get("api_key")
    )


def _mime_type_for_extension(file_ext: str) -> str:
    if file_ext == "docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if file_ext == "xlsx":
        return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if file_ext == "pdf":
        return "application/pdf"
    return "application/octet-stream"


def _read_sheet_rows(workbook, sheet_name: str, expected_headers: list[str], entity_kind: str) -> list[dict]:
    sheet = workbook[sheet_name]
    headers = [_normalize_header(cell.value) for cell in sheet[1]]
    missing_headers = [header for header in expected_headers if header not in headers]
    if missing_headers:
        raise ValueError(f"Sheet '{sheet_name}' is missing required headers: " + ", ".join(missing_headers))

    rows = []
    for source_row_number, row_cells in enumerate(sheet.iter_rows(min_row=2), start=2):
        values: dict[str, Any] = {}
        for index, header in enumerate(headers):
            if not header:
                continue
            values[header] = row_cells[index].value
        if _row_is_empty(values):
            continue
        rows.append(
            {
                "id": None,
                "sheet_name": sheet_name,
                "source_row_number": source_row_number,
                "entity_kind": entity_kind,
                "values": values,
            }
        )
    return rows


def _validate_standard_row(row: dict, existing_by_code: dict[str, dict]) -> dict:
    values = row["values"]
    issues = []
    code = _normalize_code(values.get("code"))
    name = _trim_text(values.get("name"))
    status = _normalize_status(values.get("status"), VALID_STANDARD_STATUSES)
    metadata, metadata_issue = _parse_json_object(values.get("metadata_json"), "metadata_json")
    existing = existing_by_code.get(_code_key(code)) if code else None

    if not code:
        issues.append(_issue("required", "code", "标准编码不能为空"))
    if not name:
        issues.append(_issue("required", "name", "标准名称不能为空"))
    if status is None:
        issues.append(_issue("invalid_status", "status", "标准状态仅支持 draft / active / archived"))
    if metadata_issue:
        issues.append(metadata_issue)

    normalized = {
        "code": code,
        "name": name,
        "version_label": _blank_to_none(values.get("version_label")),
        "thumbnail_url": _blank_to_none(values.get("thumbnail_url")),
        "status": status or "active",
        "metadata": metadata,
        "existing_standard_id": str(existing["id"]) if existing else None,
    }
    return _validated_row(row, normalized, issues, conflict=existing is not None)


def _validate_pbs_level_row(row: dict, context: dict) -> dict:
    values = row["values"]
    issues = []
    level_no = _int_or_none(values.get("level_no"))
    code = _normalize_code(values.get("code"))
    name = _trim_text(values.get("name"))

    if level_no is None or level_no <= 0:
        issues.append(_issue("invalid_level_no", "level_no", "PBS 层级编号必须是正整数"))
    elif context["pbs_level_counts"][level_no] > 1:
        issues.append(_issue("duplicate_level_no_in_file", "level_no", "Excel 文件内存在重复 PBS 层级编号"))
    if not code:
        issues.append(_issue("required", "code", "PBS 层级编码不能为空"))
    elif context["pbs_code_counts"][_code_key(code)] > 1:
        issues.append(_issue("duplicate_code_in_file", "code", "Excel 文件内存在重复 PBS 层级编码"))
    if not name:
        issues.append(_issue("required", "name", "PBS 层级名称不能为空"))

    normalized = {"level_no": level_no or 1, "code": code, "name": name, "description": _blank_to_none(values.get("description"))}
    conflict = context["target_mode"] == "merge" and _code_key(code) in context["target_existing"]["pbs_level"]
    return _validated_row(row, normalized, issues, conflict=conflict)


def _validate_class_row(row: dict, context: dict, *, kind: Literal["tag_class", "equipment_class", "document_type"]) -> dict:
    values = row["values"]
    issues = []
    code = _normalize_code(values.get("code"))
    name = _trim_text(values.get("name"))
    parent_code = _normalize_code(values.get("parent_code"))
    status = _normalize_status(values.get("status"), VALID_DEFINITION_STATUSES)
    if kind == "tag_class":
        code_counts = context["tag_class_code_counts"]
        known_codes = context["tag_class_codes"]
        level_map = context["tag_class_levels"]
    elif kind == "equipment_class":
        code_counts = context["equipment_class_code_counts"]
        known_codes = context["equipment_class_codes"]
        level_map = context["equipment_class_levels"]
    else:
        code_counts = context["document_type_code_counts"]
        known_codes = context["document_type_codes"]
        level_map = context["document_type_levels"]
    level_no = _int_or_none(values.get("level_no")) or level_map.get(_code_key(code), 1)

    if not code:
        issues.append(_issue("required", "code", "类型编码不能为空"))
    elif code_counts[_code_key(code)] > 1:
        issues.append(_issue("duplicate_code_in_file", "code", "Excel 文件内存在重复类型编码"))
    if not name:
        issues.append(_issue("required", "name", "类型名称不能为空"))
    if parent_code and _code_key(parent_code) not in known_codes:
        issues.append(_issue("parent_not_found", "parent_code", "父级编码在当前 Excel 中不存在"))
    if level_no <= 0:
        issues.append(_issue("invalid_level_no", "level_no", "层级编号必须是正整数"))
    if status is None:
        issues.append(_issue("invalid_status", "status", "状态仅支持 draft / active / deprecated / archived"))

    normalized = {
        "code": code,
        "name": name,
        "parent_code": parent_code or None,
        "level_no": level_no,
        "description": _blank_to_none(values.get("description")),
        "status": status or "active",
    }
    conflict = context["target_mode"] == "merge" and _code_key(code) in context["target_existing"][kind]
    return _validated_row(row, normalized, issues, conflict=conflict)


def _validate_document_type_row(row: dict, context: dict) -> dict:
    validated = _validate_class_row(row, context, kind="document_type")
    metadata, metadata_issue = _parse_json_object(row["values"].get("metadata_json"), "metadata_json")
    if metadata_issue:
        validated["issues"].append(metadata_issue)
        validated["status"] = "error"
    validated["normalized_values"] = {
        **validated["normalized_values"],
        "allowed_extensions": _parse_list(row["values"].get("allowed_extensions")),
        "metadata": metadata,
    }
    return validated


def _validate_attribute_row(row: dict, context: dict, *, kind: Literal["tag_attribute", "equipment_attribute", "document_attribute"]) -> dict:
    values = row["values"]
    issues = []
    owner_field = "owner_class_code" if kind == "tag_attribute" else "owner_document_type_code"
    if kind == "equipment_attribute":
        owner_field = "owner_class_code"
    owner_code = _normalize_code(values.get(owner_field))
    code = _normalize_code(values.get("code"))
    name = _trim_text(values.get("name"))
    value_type = _trim_text(values.get("value_type")).lower() or "string"
    is_required = _parse_bool(values.get("is_required"))
    sort_order = _int_or_none(values.get("sort_order"))
    status = _normalize_status(values.get("status"), VALID_DEFINITION_STATUSES)
    if kind == "tag_attribute":
        owner_codes = context["tag_class_codes"]
        counts = context["tag_attribute_counts"]
    elif kind == "equipment_attribute":
        owner_codes = context["equipment_class_codes"]
        counts = context["equipment_attribute_counts"]
    else:
        owner_codes = context["document_type_codes"]
        counts = context["document_attribute_counts"]
    owner_key = _code_key(owner_code) or "__common__"

    if owner_code and _code_key(owner_code) not in owner_codes:
        issues.append(_issue("owner_not_found", owner_field, "属性归属类型编码在当前 Excel 中不存在"))
    if not code:
        issues.append(_issue("required", "code", "属性编码不能为空"))
    elif counts[(owner_key, _code_key(code))] > 1:
        issues.append(_issue("duplicate_code_in_file", "code", "同一归属下存在重复属性编码"))
    if not name:
        issues.append(_issue("required", "name", "属性名称不能为空"))
    if value_type not in VALID_VALUE_TYPES:
        issues.append(_issue("invalid_value_type", "value_type", "属性数据类型不受支持"))
    if is_required is None:
        issues.append(_issue("invalid_boolean", "is_required", "是否必填必须是 true/false、是/否 或 1/0"))
    if sort_order is not None and sort_order < 0:
        issues.append(_issue("invalid_sort_order", "sort_order", "排序号不能小于 0"))
    if status is None:
        issues.append(_issue("invalid_status", "status", "状态仅支持 draft / active / deprecated / archived"))

    normalized = {
        owner_field: owner_code or None,
        "group_name": _blank_to_none(values.get("group_name")),
        "code": code,
        "name": name,
        "value_type": value_type,
        "is_required": bool(is_required),
        "unit_family": _blank_to_none(values.get("unit_family")),
        "enum_options": _parse_list(values.get("enum_options")),
        "description": _blank_to_none(values.get("description")),
        "sort_order": sort_order or 0,
        "status": status or "active",
    }
    conflict = (
        context["target_mode"] == "merge"
        and (owner_key, _code_key(code)) in context["target_existing"][kind]
    )
    return _validated_row(row, normalized, issues, conflict=conflict)


def _validate_tag_equipment_relationship_row(row: dict, context: dict) -> dict:
    values = row["values"]
    issues = []
    tag_class_code = _normalize_code(values.get("tag_class_code"))
    equipment_class_code = _normalize_code(values.get("equipment_class_code"))
    status = _normalize_status(values.get("status"), VALID_DEFINITION_STATUSES)
    pair_key = (_code_key(tag_class_code), _code_key(equipment_class_code))

    if not tag_class_code:
        issues.append(_issue("required", "tag_class_code", "Tag Class 编码不能为空"))
    elif _code_key(tag_class_code) not in context["tag_class_codes"]:
        issues.append(_issue("tag_class_not_found", "tag_class_code", "Tag Class 编码在当前 Excel 中不存在"))
    if not equipment_class_code:
        issues.append(_issue("required", "equipment_class_code", "Equipment Class 编码不能为空"))
    elif _code_key(equipment_class_code) not in context["equipment_class_codes"]:
        issues.append(_issue("equipment_class_not_found", "equipment_class_code", "Equipment Class 编码在当前 Excel 中不存在"))
    if context["tag_equipment_relationship_counts"][pair_key] > 1:
        issues.append(_issue("duplicate_relationship_in_file", "equipment_class_code", "Excel 文件内存在重复 Tag/Equipment Class 映射"))
    if status is None:
        issues.append(_issue("invalid_status", "status", "状态仅支持 draft / active / deprecated / archived"))

    normalized = {
        "tag_class_code": tag_class_code,
        "equipment_class_code": equipment_class_code,
        "reason": _blank_to_none(values.get("reason")),
        "status": status or "active",
    }
    conflict = (
        context["target_mode"] == "merge"
        and pair_key in context["target_existing"]["tag_equipment_class_relationship"]
    )
    return _validated_row(row, normalized, issues, conflict=conflict)


def _document_type_reference_exists(document_type_code: str | None, cfihos_code: str | None, context: dict) -> bool:
    if document_type_code and _code_key(document_type_code) in context["document_type_codes"]:
        return True
    if cfihos_code and _code_key(cfihos_code) in context["document_type_cfihos_codes"]:
        return True
    return False


def _class_reference_exists(class_code: str | None, asset_scope: str | None, context: dict) -> bool:
    if not class_code:
        return False
    class_key = _code_key(class_code)
    scope = _normalize_context_code(asset_scope, default="") or ""
    if scope == "tag":
        return class_key in context["tag_class_codes"] or class_key in context["equipment_class_codes"]
    if scope in {"equipment", "model_part"}:
        return class_key in context["equipment_class_codes"] or class_key in context["tag_class_codes"]
    return class_key in context["tag_class_codes"] or class_key in context["equipment_class_codes"]


def _validate_discipline_row(row: dict, context: dict) -> dict:
    values = row["values"]
    issues = []
    cfihos_unique_code = _normalize_code(values.get("cfihos_unique_code"))
    code = _normalize_code(values.get("code"))
    name = _trim_text(values.get("name"))
    status = _normalize_status(values.get("status"), VALID_RULE_STATUSES)
    metadata, metadata_issue = _parse_json_object(values.get("metadata_json"), "metadata_json")

    if not code:
        issues.append(_issue("required", "code", "Discipline 编码不能为空"))
    elif context["discipline_code_counts"][_code_key(code)] > 1:
        issues.append(_issue("duplicate_code_in_file", "code", "Excel 文件内存在重复 Discipline 编码"))
    if cfihos_unique_code and context["discipline_cfihos_code_counts"][_code_key(cfihos_unique_code)] > 1:
        issues.append(_issue("duplicate_cfihos_code_in_file", "cfihos_unique_code", "Excel 文件内存在重复 Discipline CFIHOS unique code"))
    if not name:
        issues.append(_issue("required", "name", "Discipline 名称不能为空"))
    if status is None:
        issues.append(_issue("invalid_status", "status", "状态仅支持 active / deprecated / archived"))
    if metadata_issue:
        issues.append(metadata_issue)

    normalized = {
        "cfihos_unique_code": cfihos_unique_code or None,
        "code": code,
        "name": name,
        "description": _blank_to_none(values.get("description")),
        "status": status or "active",
        "metadata": metadata,
    }
    conflict = (
        context["target_mode"] == "merge"
        and (
            _code_key(code) in context["target_existing"]["discipline"]
            or (cfihos_unique_code and _code_key(cfihos_unique_code) in context["target_existing"]["discipline_cfihos"])
        )
    )
    return _validated_row(row, normalized, issues, conflict=conflict)


def _validate_discipline_document_type_row(row: dict, context: dict) -> dict:
    values = row["values"]
    issues = []
    cfihos_unique_code = _normalize_code(values.get("cfihos_unique_code"))
    discipline_cfihos_code = _normalize_code(values.get("discipline_cfihos_unique_code"))
    discipline_code = _normalize_code(values.get("discipline_code"))
    document_type_cfihos_code = _normalize_code(values.get("document_type_cfihos_unique_code"))
    document_type_code = _normalize_code(values.get("document_type_code")) or context["document_type_code_by_cfihos"].get(_code_key(document_type_cfihos_code))
    asset_scope = _normalize_context_code(values.get("asset_scope"))
    perspective = _normalize_context_code(values.get("perspective"), default="standard") or "standard"
    lifecycle_phase = _normalize_context_code(values.get("lifecycle_phase"), default="unspecified") or "unspecified"
    status = _normalize_status(values.get("status"), VALID_RULE_STATUSES)
    metadata, metadata_issue = _parse_json_object(values.get("metadata_json"), "metadata_json")
    relationship_key = _discipline_document_type_key(
        {
            **values,
            "discipline_code": discipline_code,
            "discipline_cfihos_unique_code": discipline_cfihos_code,
            "document_type_code": document_type_code,
            "document_type_cfihos_unique_code": document_type_cfihos_code,
            "asset_scope": asset_scope,
            "perspective": perspective,
            "lifecycle_phase": lifecycle_phase,
        }
    )

    if not discipline_code and not discipline_cfihos_code:
        issues.append(_issue("required", "discipline_code", "Discipline 引用不能为空"))
    elif not (
        (discipline_code and _code_key(discipline_code) in context["discipline_codes"])
        or (discipline_cfihos_code and _code_key(discipline_cfihos_code) in context["discipline_cfihos_codes"])
    ):
        issues.append(_issue("discipline_not_found", "discipline_code", "Discipline 引用在当前 Excel 中不存在"))
    if not document_type_code and not document_type_cfihos_code:
        issues.append(_issue("required", "document_type_code", "Document Type 引用不能为空"))
    elif not _document_type_reference_exists(document_type_code, document_type_cfihos_code, context):
        issues.append(_issue("document_type_not_found", "document_type_code", "Document Type 引用在当前 Excel 中不存在"))
    if context["discipline_document_type_counts"][relationship_key] > 1:
        issues.append(_issue("duplicate_relationship_in_file", "document_type_code", "Excel 文件内存在重复 Discipline/Document Type 规则"))
    if status is None:
        issues.append(_issue("invalid_status", "status", "状态仅支持 active / deprecated / archived"))
    if metadata_issue:
        issues.append(metadata_issue)

    normalized = {
        "cfihos_unique_code": cfihos_unique_code or None,
        "discipline_cfihos_unique_code": discipline_cfihos_code or None,
        "discipline_code": discipline_code or None,
        "document_type_cfihos_unique_code": document_type_cfihos_code or None,
        "document_type_code": document_type_code or None,
        "document_type_name": _blank_to_none(values.get("document_type_name")),
        "short_code": _blank_to_none(values.get("short_code")),
        "asset_scope": asset_scope,
        "representation_type": _blank_to_none(values.get("representation_type")),
        "native_file_delivery_timing": _blank_to_none(values.get("native_file_delivery_timing")),
        "perspective": perspective,
        "lifecycle_phase": lifecycle_phase,
        "status": status or "active",
        "metadata": metadata,
    }
    conflict = context["target_mode"] == "merge" and relationship_key in context["target_existing"]["discipline_document_type"]
    return _validated_row(row, normalized, issues, conflict=conflict)


def _validate_class_document_requirement_row(row: dict, context: dict) -> dict:
    values = row["values"]
    issues = []
    cfihos_unique_code = _normalize_code(values.get("cfihos_unique_code"))
    class_code = _normalize_code(values.get("class_code"))
    asset_scope = _normalize_context_code(values.get("asset_scope"))
    document_type_cfihos_code = _normalize_code(values.get("document_type_cfihos_unique_code"))
    document_type_code = _normalize_code(values.get("document_type_code")) or context["document_type_code_by_cfihos"].get(_code_key(document_type_cfihos_code))
    perspective = _normalize_context_code(values.get("perspective"), default="standard") or "standard"
    lifecycle_phase = _normalize_context_code(values.get("lifecycle_phase"), default="unspecified") or "unspecified"
    status = _normalize_status(values.get("status"), VALID_RULE_STATUSES)
    metadata, metadata_issue = _parse_json_object(values.get("metadata_json"), "metadata_json")
    requirement_key = _class_document_requirement_key(
        {
            **values,
            "class_code": class_code,
            "asset_scope": asset_scope,
            "document_type_code": document_type_code,
            "document_type_cfihos_unique_code": document_type_cfihos_code,
            "perspective": perspective,
            "lifecycle_phase": lifecycle_phase,
        }
    )

    if not class_code:
        issues.append(_issue("required", "class_code", "Tag/Equipment Class 引用不能为空"))
    elif not _class_reference_exists(class_code, asset_scope, context):
        issues.append(_issue("class_not_found", "class_code", "Tag/Equipment Class 引用在当前 Excel 中不存在"))
    if not document_type_code and not document_type_cfihos_code:
        issues.append(_issue("required", "document_type_code", "Document Type 引用不能为空"))
    elif not _document_type_reference_exists(document_type_code, document_type_cfihos_code, context):
        issues.append(_issue("document_type_not_found", "document_type_code", "Document Type 引用在当前 Excel 中不存在"))
    if context["class_document_requirement_counts"][requirement_key] > 1:
        issues.append(_issue("duplicate_requirement_in_file", "document_type_code", "Excel 文件内存在重复 Class/Document Type 要求"))
    if status is None:
        issues.append(_issue("invalid_status", "status", "状态仅支持 active / deprecated / archived"))
    if metadata_issue:
        issues.append(metadata_issue)

    normalized = {
        "cfihos_unique_code": cfihos_unique_code or None,
        "class_code": class_code,
        "class_name": _blank_to_none(values.get("class_name")),
        "asset_scope": asset_scope,
        "source_standard_cfihos_code": _blank_to_none(values.get("source_standard_cfihos_code")),
        "source_standard_code": _blank_to_none(values.get("source_standard_code")),
        "document_type_cfihos_unique_code": document_type_cfihos_code or None,
        "document_type_code": document_type_code or None,
        "document_type_name": _blank_to_none(values.get("document_type_name")),
        "perspective": perspective,
        "lifecycle_phase": lifecycle_phase,
        "status": status or "active",
        "metadata": metadata,
    }
    conflict = context["target_mode"] == "merge" and requirement_key in context["target_existing"]["class_document_requirement"]
    return _validated_row(row, normalized, issues, conflict=conflict)


def _store_standard_import_job(
    *,
    filename: str,
    file_ext: str,
    file_size: int,
    checksum_sha256: str,
    target_mode: str,
    target_standard_id: str | None,
    validated: dict,
    chunks: list[dict],
) -> dict:
    standard = validated["standard"]
    summary = validated["summary"]
    with get_connection() as connection:
        with connection.cursor() as cursor:
            cursor.execute(
                """
                INSERT INTO standard_import_job (
                    filename, file_ext, file_size, checksum_sha256,
                    target_mode, target_standard_id, summary, source_standard_code
                )
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING id
                """,
                (
                    filename,
                    file_ext,
                    file_size,
                    checksum_sha256,
                    target_mode,
                    target_standard_id,
                    Json(summary),
                    standard.get("code"),
                ),
            )
            created = cursor.fetchone()
            file_id = _insert_import_file(
                cursor,
                str(created["id"]),
                filename=filename,
                file_ext=file_ext,
                file_size=file_size,
                checksum_sha256=checksum_sha256,
            )
            _insert_import_chunks(cursor, str(created["id"]), file_id, chunks)
            _insert_import_rows(cursor, str(created["id"]), validated["rows"])
        connection.commit()
    return created


def _insert_import_file(
    cursor,
    job_id: str,
    *,
    filename: str,
    file_ext: str,
    file_size: int,
    checksum_sha256: str,
) -> str:
    cursor.execute(
        """
        INSERT INTO standard_import_file (
            job_id, original_filename, file_ext, mime_type, size_bytes,
            checksum_sha256, parser_profile, metadata
        )
        VALUES (%s, %s, %s, %s, %s, %s, %s, '{}'::jsonb)
        RETURNING id
        """,
        (
            job_id,
            filename,
            file_ext,
            _mime_type_for_extension(file_ext),
            file_size,
            checksum_sha256,
            "table-first-v1",
        ),
    )
    return str(cursor.fetchone()["id"])


def _insert_import_chunks(cursor, job_id: str, file_id: str, chunks: list[dict]) -> None:
    for index, chunk in enumerate(chunks, start=1):
        cursor.execute(
            """
            INSERT INTO standard_import_chunk (
                job_id, file_id, chunk_no, source_kind, sheet_name, page_no,
                table_index, heading_path, content, metadata
            )
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """,
            (
                job_id,
                file_id,
                chunk.get("chunk_no") or index,
                chunk.get("source_kind", "table"),
                chunk.get("sheet_name"),
                chunk.get("page_no"),
                chunk.get("table_index"),
                chunk.get("heading_path") or [],
                chunk.get("content") or "",
                Json(chunk.get("metadata") or {}),
            ),
        )


def _insert_import_rows(cursor, job_id: str, rows: list[dict]) -> None:
    for row in rows:
        cursor.execute(
            """
            INSERT INTO standard_import_row (
                job_id, row_number, source_kind, sheet_name, page_no, table_index,
                source_row_number, entity_kind, values, normalized_values, issues,
                status, action, confidence, evidence
            )
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """,
            (
                job_id,
                row["row_number"],
                row.get("source_kind", "template"),
                row["sheet_name"],
                row.get("page_no"),
                row.get("table_index"),
                row["source_row_number"],
                row["entity_kind"],
                Json(row["values"]),
                Json(row["normalized_values"]),
                Json(row["issues"]),
                row["status"],
                row.get("action"),
                row.get("confidence", 0.65),
                Json(row.get("evidence") or []),
            ),
        )


def _replace_standard_import_rows(job_id: str, rows: list[dict], *, action_by_row: dict[str, str | None]) -> None:
    job = _ensure_standard_import_job_exists(job_id)
    target_bundle = load_standard_export_bundle(str(job["target_standard_id"])) if job.get("target_standard_id") else None
    row_values = []
    for row in rows:
        row_values.append(
            {
                "id": row["id"],
                "row_number": row["row_number"],
                "source_kind": row.get("source_kind", "table"),
                "sheet_name": row.get("sheet_name"),
                "page_no": row.get("page_no"),
                "table_index": row.get("table_index"),
                "source_row_number": row["source_row_number"],
                "entity_kind": row["entity_kind"],
                "values": row["values"],
                "confidence": row.get("confidence", 0.65),
                "evidence": row.get("evidence") or [],
                "action": action_by_row.get(row["id"]),
            }
        )
    validated = _validate_extracted_standard_rows(
        row_values,
        existing_standards=fetch_all("SELECT id, code FROM standard"),
        target_mode=job.get("target_mode") or "new",
        target_standard_bundle=target_bundle,
    )
    next_rows = _apply_row_actions(validated["rows"], action_by_row)
    summary = _summarize_rows(next_rows)

    with get_connection() as connection:
        with connection.cursor() as cursor:
            cursor.execute(
                """
                UPDATE standard_import_job
                SET summary = %s,
                    source_standard_code = %s,
                    updated_at = now()
                WHERE id = %s
                """,
                (Json(summary), validated["standard"].get("code"), job_id),
            )
            for row in next_rows:
                cursor.execute(
                    """
                    UPDATE standard_import_row
                    SET values = %s,
                        normalized_values = %s,
                        issues = %s,
                        status = %s,
                        action = %s,
                        confidence = %s,
                        evidence = %s,
                        updated_at = now()
                    WHERE id = %s
                      AND job_id = %s
                    """,
                    (
                        Json(row["values"]),
                        Json(row["normalized_values"]),
                        Json(row["issues"]),
                        row["status"],
                        row.get("action"),
                        row.get("confidence", 0.65),
                        Json(row.get("evidence") or []),
                        row["id"],
                        job_id,
                    ),
                )
        connection.commit()


def _load_standard_import_rows(job_id: str) -> list[dict]:
    rows = fetch_all(
        """
        SELECT id, row_number, source_kind, sheet_name, page_no, table_index,
               source_row_number, entity_kind, values, normalized_values, issues,
               status, action, confidence, evidence, created_at, updated_at
        FROM standard_import_row
        WHERE job_id = %s
        ORDER BY row_number
        """,
        (job_id,),
    )
    return [_serialize_import_row(row) for row in rows]


def _ensure_standard_import_job_exists(job_id: str) -> dict:
    job = fetch_one(
        """
        SELECT id, filename, file_ext, file_size, checksum_sha256, target_mode,
               summary, status, target_standard_id, source_standard_code,
               committed_at, created_at, updated_at
        FROM standard_import_job
        WHERE id = %s
        """,
        (job_id,),
    )
    if job is None:
        raise ValueError("Standard import job not found")
    return job


def _standard_code_exists(code: str) -> bool:
    return fetch_one("SELECT id FROM standard WHERE lower(code) = lower(%s)", (code,)) is not None


def _fetch_standard_by_code(cursor, code: str) -> dict | None:
    cursor.execute(
        """
        SELECT id, code
        FROM standard
        WHERE lower(code) = lower(%s)
        """,
        (code,),
    )
    return cursor.fetchone()


def _insert_standard(cursor, standard: dict, code: str) -> str:
    cursor.execute(
        """
        INSERT INTO standard (code, name, version_label, thumbnail_url, status, metadata)
        VALUES (%s, %s, %s, %s, %s, %s)
        RETURNING id
        """,
        (
            code,
            standard["name"],
            standard.get("version_label"),
            standard.get("thumbnail_url"),
            standard.get("status", "active"),
            Json(standard.get("metadata") or {}),
        ),
    )
    return str(cursor.fetchone()["id"])


def _update_standard(cursor, standard_id: str, standard: dict) -> None:
    cursor.execute(
        """
        UPDATE standard
        SET name = %s,
            version_label = %s,
            thumbnail_url = %s,
            status = %s,
            metadata = %s,
            updated_at = now()
        WHERE id = %s
        """,
        (
            standard["name"],
            standard.get("version_label"),
            standard.get("thumbnail_url"),
            standard.get("status", "active"),
            Json(standard.get("metadata") or {}),
            standard_id,
        ),
    )


def _upsert_pbs_levels(cursor, standard_id: str, rows: list[dict]) -> tuple[int, int]:
    created = 0
    updated = 0
    for row in rows:
        cursor.execute(
            """
            SELECT id
            FROM pbs_level_template
            WHERE standard_id = %s
              AND lower(code) = lower(%s)
            """,
            (standard_id, row["code"]),
        )
        existing = cursor.fetchone()
        if existing:
            cursor.execute(
                """
                UPDATE pbs_level_template
                SET level_no = %s,
                    code = %s,
                    name = %s,
                    description = %s,
                    updated_at = now()
                WHERE id = %s
                """,
                (row["level_no"], row["code"], row["name"], row.get("description"), existing["id"]),
            )
            updated += 1
        else:
            cursor.execute(
                """
                INSERT INTO pbs_level_template (standard_id, level_no, code, name, description)
                VALUES (%s, %s, %s, %s, %s)
                """,
                (standard_id, row["level_no"], row["code"], row["name"], row.get("description")),
            )
            created += 1
    return created, updated


def _upsert_classes(cursor, standard_id: str, rows: list[dict], *, applies_to: Literal["tag", "document", "equipment"]) -> dict[str, dict]:
    class_ids: dict[str, dict] = {}
    for row in rows:
        metadata = _class_metadata(row) if applies_to == "document" else {}
        cursor.execute(
            """
            SELECT id
            FROM class
            WHERE standard_id = %s
              AND lower(code) = lower(%s)
              AND (applies_to = %s OR (%s IN ('tag', 'document') AND applies_to = 'both'))
            ORDER BY CASE WHEN applies_to = %s THEN 0 ELSE 1 END
            LIMIT 1
            """,
            (standard_id, row["code"], applies_to, applies_to, applies_to),
        )
        existing = cursor.fetchone()
        if existing:
            cursor.execute(
                """
                UPDATE class
                SET code = %s,
                    name = %s,
                    level_no = %s,
                    description = %s,
                    status = %s,
                    metadata = %s,
                    updated_at = now()
                WHERE id = %s
                """,
                (
                    row["code"],
                    row["name"],
                    row.get("level_no", 1),
                    row.get("description"),
                    row.get("status", "active"),
                    Json(metadata),
                    existing["id"],
                ),
            )
            class_ids[_code_key(row["code"])] = {"id": str(existing["id"]), "created": False}
        else:
            cursor.execute(
                """
                INSERT INTO class (standard_id, code, name, parent_id, level_no, description, status, metadata, applies_to)
                VALUES (%s, %s, %s, NULL, %s, %s, %s, %s, %s)
                RETURNING id
                """,
                (
                    standard_id,
                    row["code"],
                    row["name"],
                    row.get("level_no", 1),
                    row.get("description"),
                    row.get("status", "active"),
                    Json(metadata),
                    applies_to,
                ),
            )
            class_ids[_code_key(row["code"])] = {"id": str(cursor.fetchone()["id"]), "created": True}

    for row in rows:
        item = class_ids[_code_key(row["code"])]
        parent_id = None
        if row.get("parent_code"):
            parent = class_ids.get(_code_key(row["parent_code"]))
            parent_id = parent["id"] if parent else None
        cursor.execute(
            """
            UPDATE class
            SET parent_id = %s,
                level_no = %s,
                updated_at = now()
            WHERE id = %s
            """,
            (parent_id, row.get("level_no", 1), item["id"]),
        )
    return class_ids


def _upsert_attributes(
    cursor,
    standard_id: str,
    rows: list[dict],
    *,
    applies_to: Literal["tag", "document", "equipment"],
    owner_ids: dict[str, dict],
) -> tuple[int, int]:
    created = 0
    updated = 0
    owner_field = "owner_document_type_code" if applies_to == "document" else "owner_class_code"
    for row in rows:
        owner_code = row.get(owner_field)
        if owner_code and _code_key(owner_code) not in owner_ids:
            continue
        owner_id = owner_ids[_code_key(owner_code)]["id"] if owner_code else None
        standard_owner_id = None if owner_id else standard_id
        if owner_id:
            cursor.execute(
                """
                SELECT id
                FROM attribute_definition
                WHERE class_id = %s
                  AND lower(code) = lower(%s)
                  AND (applies_to = %s OR (%s IN ('tag', 'document') AND applies_to = 'both'))
                LIMIT 1
                """,
                (owner_id, row["code"], applies_to, applies_to),
            )
        else:
            cursor.execute(
                """
                SELECT id
                FROM attribute_definition
                WHERE class_id IS NULL
                  AND standard_id = %s
                  AND lower(code) = lower(%s)
                  AND (applies_to = %s OR (%s IN ('tag', 'document') AND applies_to = 'both'))
                LIMIT 1
                """,
                (standard_id, row["code"], applies_to, applies_to),
            )
        existing = cursor.fetchone()
        params = (
            row.get("group_name"),
            row["code"],
            row["name"],
            row["value_type"],
            row["is_required"],
            row.get("unit_family"),
            Json(row.get("enum_options") or []),
            row.get("description"),
            row.get("sort_order", 0),
            row.get("status", "active"),
        )
        if existing:
            cursor.execute(
                """
                UPDATE attribute_definition
                SET group_name = %s,
                    code = %s,
                    name = %s,
                    value_type = %s,
                    is_required = %s,
                    unit_family = %s,
                    enum_options = %s,
                    description = %s,
                    sort_order = %s,
                    status = %s,
                    updated_at = now()
                WHERE id = %s
                """,
                (*params, existing["id"]),
            )
            updated += 1
        else:
            cursor.execute(
                """
                INSERT INTO attribute_definition (
                    class_id, standard_id, group_name, code, name, value_type, is_required,
                    unit_family, enum_options, description, sort_order, status, metadata, applies_to
                )
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, '{}'::jsonb, %s)
                """,
                (owner_id, standard_owner_id, *params, applies_to),
            )
            created += 1
    return created, updated


def _upsert_tag_equipment_class_relationships(
    cursor,
    standard_id: str,
    rows: list[dict],
    *,
    tag_class_ids: dict[str, dict],
    equipment_class_refs: dict[str, dict],
) -> tuple[int, int]:
    created = 0
    updated = 0
    for row in rows:
        tag_class = tag_class_ids.get(_code_key(row.get("tag_class_code")))
        equipment_class = equipment_class_refs.get(_code_key(row.get("equipment_class_code")))
        if not tag_class or not equipment_class:
            continue
        cursor.execute(
            """
            SELECT id
            FROM class_relationship
            WHERE standard_id = %s
              AND source_class_id = %s
              AND target_class_id = %s
              AND relationship_type = %s
            LIMIT 1
            """,
            (standard_id, tag_class["id"], equipment_class["id"], TAG_EQUIPMENT_CLASS_RELATIONSHIP),
        )
        existing = cursor.fetchone()
        if existing:
            cursor.execute(
                """
                UPDATE class_relationship
                SET reason = %s,
                    status = %s,
                    updated_at = now()
                WHERE id = %s
                """,
                (row.get("reason"), row.get("status", "active"), existing["id"]),
            )
            updated += 1
        else:
            cursor.execute(
                """
                INSERT INTO class_relationship (
                    standard_id, source_class_id, target_class_id, relationship_type, reason, status, metadata
                )
                VALUES (%s, %s, %s, %s, %s, %s, '{}'::jsonb)
                """,
                (
                    standard_id,
                    tag_class["id"],
                    equipment_class["id"],
                    TAG_EQUIPMENT_CLASS_RELATIONSHIP,
                    row.get("reason"),
                    row.get("status", "active"),
                ),
            )
            created += 1
    return created, updated


def _upsert_disciplines(cursor, standard_id: str, rows: list[dict]) -> dict[str, Any]:
    refs: dict[str, Any] = {
        "by_code": {},
        "by_cfihos": {},
        "created_count": 0,
        "updated_count": 0,
    }
    for row in rows:
        cursor.execute(
            """
            SELECT id
            FROM discipline
            WHERE standard_id = %s
              AND status <> 'archived'
              AND (
                    lower(code) = lower(%s)
                    OR (%s::text IS NOT NULL AND lower(cfihos_unique_code) = lower(%s::text))
                  )
            ORDER BY CASE WHEN lower(code) = lower(%s) THEN 0 ELSE 1 END
            LIMIT 1
            """,
            (
                standard_id,
                row["code"],
                row.get("cfihos_unique_code"),
                row.get("cfihos_unique_code"),
                row["code"],
            ),
        )
        existing = cursor.fetchone()
        if existing:
            cursor.execute(
                """
                UPDATE discipline
                SET cfihos_unique_code = %s,
                    code = %s,
                    name = %s,
                    description = %s,
                    status = %s,
                    metadata = %s,
                    updated_at = now()
                WHERE id = %s
                """,
                (
                    row.get("cfihos_unique_code"),
                    row["code"],
                    row["name"],
                    row.get("description"),
                    row.get("status", "active"),
                    Json(row.get("metadata") or {}),
                    existing["id"],
                ),
            )
            item = {"id": str(existing["id"]), "created": False}
            refs["updated_count"] += 1
        else:
            cursor.execute(
                """
                INSERT INTO discipline (
                    standard_id, cfihos_unique_code, code, name, description, status, metadata
                )
                VALUES (%s, %s, %s, %s, %s, %s, %s)
                RETURNING id
                """,
                (
                    standard_id,
                    row.get("cfihos_unique_code"),
                    row["code"],
                    row["name"],
                    row.get("description"),
                    row.get("status", "active"),
                    Json(row.get("metadata") or {}),
                ),
            )
            item = {"id": str(cursor.fetchone()["id"]), "created": True}
            refs["created_count"] += 1

        refs["by_code"][_code_key(row["code"])] = item
        if row.get("cfihos_unique_code"):
            refs["by_cfihos"][_code_key(row.get("cfihos_unique_code"))] = item
    return refs


def _document_type_cfihos_refs(document_type_ids: dict[str, dict], rows: list[dict]) -> dict[str, dict]:
    refs: dict[str, dict] = {}
    for row in rows:
        item = document_type_ids.get(_code_key(row.get("code")))
        metadata = row.get("metadata") or {}
        cfihos_code = _code_key(metadata.get("cfihos_unique_code"))
        if item and cfihos_code:
            refs[cfihos_code] = item
    return refs


def _resolve_discipline_ref(row: dict, discipline_refs: dict[str, Any]) -> dict | None:
    cfihos_code = _code_key(row.get("discipline_cfihos_unique_code"))
    if cfihos_code:
        found = discipline_refs["by_cfihos"].get(cfihos_code)
        if found:
            return found
    return discipline_refs["by_code"].get(_code_key(row.get("discipline_code")))


def _resolve_document_type_ref(
    row: dict,
    *,
    document_type_ids: dict[str, dict],
    document_type_cfihos_refs: dict[str, dict],
) -> dict | None:
    cfihos_code = _code_key(row.get("document_type_cfihos_unique_code"))
    if cfihos_code:
        found = document_type_cfihos_refs.get(cfihos_code)
        if found:
            return found
    return document_type_ids.get(_code_key(row.get("document_type_code")))


def _resolve_class_ref(row: dict, *, tag_class_ids: dict[str, dict], equipment_class_refs: dict[str, dict]) -> dict | None:
    class_code = _code_key(row.get("class_code"))
    asset_scope = _normalize_context_code(row.get("asset_scope"), default="") or ""
    if asset_scope == "tag":
        return tag_class_ids.get(class_code) or equipment_class_refs.get(class_code)
    if asset_scope in {"equipment", "model_part"}:
        return equipment_class_refs.get(class_code) or tag_class_ids.get(class_code)
    return tag_class_ids.get(class_code) or equipment_class_refs.get(class_code)


def _upsert_discipline_document_types(
    cursor,
    standard_id: str,
    rows: list[dict],
    *,
    discipline_refs: dict[str, Any],
    document_type_ids: dict[str, dict],
    document_type_cfihos_refs: dict[str, dict],
) -> tuple[int, int]:
    created = 0
    updated = 0
    for row in rows:
        discipline = _resolve_discipline_ref(row, discipline_refs)
        document_type = _resolve_document_type_ref(
            row,
            document_type_ids=document_type_ids,
            document_type_cfihos_refs=document_type_cfihos_refs,
        )
        if not discipline or not document_type:
            continue

        asset_scope_key = _normalize_context_code(row.get("asset_scope"), default="") or ""
        representation_type_key = _relationship_text_key(row.get("representation_type"))
        delivery_timing_key = _relationship_text_key(row.get("native_file_delivery_timing"))
        perspective_key = _normalize_context_code(row.get("perspective"), default="standard") or "standard"
        lifecycle_phase_key = _normalize_context_code(row.get("lifecycle_phase"), default="unspecified") or "unspecified"
        cursor.execute(
            """
            SELECT id
            FROM discipline_document_type
            WHERE standard_id = %s
              AND status <> 'archived'
              AND (
                    (%s::text IS NOT NULL AND lower(cfihos_unique_code) = lower(%s::text))
                    OR (
                        discipline_id = %s
                        AND document_type_id = %s
                        AND COALESCE(lower(asset_scope), '') = %s
                        AND COALESCE(lower(representation_type), '') = %s
                        AND COALESCE(lower(native_file_delivery_timing), '') = %s
                        AND lower(perspective) = %s
                        AND lower(lifecycle_phase) = %s
                    )
                  )
            LIMIT 1
            """,
            (
                standard_id,
                row.get("cfihos_unique_code"),
                row.get("cfihos_unique_code"),
                discipline["id"],
                document_type["id"],
                asset_scope_key,
                representation_type_key,
                delivery_timing_key,
                perspective_key,
                lifecycle_phase_key,
            ),
        )
        existing = cursor.fetchone()
        params = (
            discipline["id"],
            document_type["id"],
            row.get("cfihos_unique_code"),
            row.get("short_code"),
            row.get("asset_scope"),
            row.get("representation_type"),
            row.get("native_file_delivery_timing"),
            row.get("perspective", "standard"),
            row.get("lifecycle_phase", "unspecified"),
            row.get("status", "active"),
            Json(row.get("metadata") or {}),
        )
        if existing:
            cursor.execute(
                """
                UPDATE discipline_document_type
                SET discipline_id = %s,
                    document_type_id = %s,
                    cfihos_unique_code = %s,
                    short_code = %s,
                    asset_scope = %s,
                    representation_type = %s,
                    native_file_delivery_timing = %s,
                    perspective = %s,
                    lifecycle_phase = %s,
                    status = %s,
                    metadata = %s,
                    updated_at = now()
                WHERE id = %s
                """,
                (*params, existing["id"]),
            )
            updated += 1
        else:
            cursor.execute(
                """
                INSERT INTO discipline_document_type (
                    standard_id, discipline_id, document_type_id, cfihos_unique_code, short_code,
                    asset_scope, representation_type, native_file_delivery_timing,
                    perspective, lifecycle_phase, status, metadata
                )
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """,
                (standard_id, *params),
            )
            created += 1
    return created, updated


def _upsert_class_document_requirements(
    cursor,
    standard_id: str,
    rows: list[dict],
    *,
    tag_class_ids: dict[str, dict],
    equipment_class_refs: dict[str, dict],
    document_type_ids: dict[str, dict],
    document_type_cfihos_refs: dict[str, dict],
) -> tuple[int, int]:
    created = 0
    updated = 0
    for row in rows:
        class_ref = _resolve_class_ref(row, tag_class_ids=tag_class_ids, equipment_class_refs=equipment_class_refs)
        document_type = _resolve_document_type_ref(
            row,
            document_type_ids=document_type_ids,
            document_type_cfihos_refs=document_type_cfihos_refs,
        )
        if not class_ref or not document_type:
            continue

        asset_scope_key = _normalize_context_code(row.get("asset_scope"), default="") or ""
        source_standard_key = _code_key(row.get("source_standard_code") or row.get("source_standard_cfihos_code"))
        perspective_key = _normalize_context_code(row.get("perspective"), default="standard") or "standard"
        lifecycle_phase_key = _normalize_context_code(row.get("lifecycle_phase"), default="unspecified") or "unspecified"
        cursor.execute(
            """
            SELECT id
            FROM class_document_requirement
            WHERE standard_id = %s
              AND status <> 'archived'
              AND (
                    (%s::text IS NOT NULL AND lower(cfihos_unique_code) = lower(%s::text))
                    OR (
                        class_id = %s
                        AND document_type_id = %s
                        AND COALESCE(lower(asset_scope), '') = %s
                        AND COALESCE(lower(source_standard_code), lower(source_standard_cfihos_code), '') = %s
                        AND lower(perspective) = %s
                        AND lower(lifecycle_phase) = %s
                    )
                  )
            LIMIT 1
            """,
            (
                standard_id,
                row.get("cfihos_unique_code"),
                row.get("cfihos_unique_code"),
                class_ref["id"],
                document_type["id"],
                asset_scope_key,
                source_standard_key,
                perspective_key,
                lifecycle_phase_key,
            ),
        )
        existing = cursor.fetchone()
        params = (
            class_ref["id"],
            document_type["id"],
            row.get("cfihos_unique_code"),
            row.get("asset_scope"),
            row.get("source_standard_cfihos_code"),
            row.get("source_standard_code"),
            row.get("perspective", "standard"),
            row.get("lifecycle_phase", "unspecified"),
            row.get("status", "active"),
            Json(row.get("metadata") or {}),
        )
        if existing:
            cursor.execute(
                """
                UPDATE class_document_requirement
                SET class_id = %s,
                    document_type_id = %s,
                    cfihos_unique_code = %s,
                    asset_scope = %s,
                    source_standard_cfihos_code = %s,
                    source_standard_code = %s,
                    perspective = %s,
                    lifecycle_phase = %s,
                    status = %s,
                    metadata = %s,
                    updated_at = now()
                WHERE id = %s
                """,
                (*params, existing["id"]),
            )
            updated += 1
        else:
            cursor.execute(
                """
                INSERT INTO class_document_requirement (
                    standard_id, class_id, document_type_id, cfihos_unique_code, asset_scope,
                    source_standard_cfihos_code, source_standard_code,
                    perspective, lifecycle_phase, status, metadata
                )
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """,
                (standard_id, *params),
            )
            created += 1
    return created, updated


def _class_metadata(row: dict) -> dict:
    metadata = dict(row.get("metadata") or {})
    metadata["document"] = {"allowed_extensions": row.get("allowed_extensions") or []}
    return metadata


def _mark_import_job_committed(cursor, job_id: str, standard_id: str, summary: dict | None) -> None:
    cursor.execute(
        """
        UPDATE standard_import_job
        SET status = 'committed',
            target_standard_id = %s,
            summary = %s,
            committed_at = now(),
            updated_at = now()
        WHERE id = %s
        """,
        (standard_id, Json(summary or {}), job_id),
    )


def _commit_result(job_id: str, standard_id: str | None, created_count: int, updated_count: int, skipped_count: int) -> dict:
    return {
        "job_id": job_id,
        "standard_id": standard_id,
        "created_count": created_count,
        "updated_count": updated_count,
        "skipped_count": skipped_count,
        "failed_count": 0,
        "failures": [],
    }


def _validated_row(row: dict, normalized_values: dict, issues: list[dict], *, conflict: bool = False) -> dict:
    status = "error" if any(issue["severity"] == "error" for issue in issues) else "ready"
    if status == "ready" and conflict:
        status = "conflict"
        issues = [
            *issues,
            _issue("target_exists", "code", "目标库已存在同编码定义，需要选择更新或跳过", severity="warning"),
        ]
    return {
        "id": row.get("id"),
        "row_number": row["row_number"],
        "sheet_name": row["sheet_name"],
        "source_kind": row.get("source_kind", "template"),
        "page_no": row.get("page_no"),
        "table_index": row.get("table_index"),
        "source_row_number": row["source_row_number"],
        "entity_kind": row["entity_kind"],
        "values": {key: _cell_value(value) for key, value in row["values"].items()},
        "normalized_values": normalized_values,
        "issues": issues,
        "status": status,
        "action": row.get("action"),
        "confidence": float(row.get("confidence", 0.65)),
        "evidence": row.get("evidence", []),
    }


def _summarize_rows(rows: list[dict]) -> dict:
    status_counts = Counter(row["status"] for row in rows)
    return {
        "total_rows": len(rows),
        "ready_rows": status_counts["ready"],
        "error_rows": status_counts["error"],
        "warning_rows": status_counts["warning"],
        "conflict_rows": status_counts["conflict"],
        "can_commit": status_counts["error"] == 0 and status_counts["conflict"] == 0,
    }


def _serialize_import_row(row: dict) -> dict:
    return {
        "id": str(row["id"]) if row.get("id") is not None else None,
        "row_number": row["row_number"],
        "sheet_name": row["sheet_name"],
        "source_kind": row.get("source_kind", "template"),
        "page_no": row.get("page_no"),
        "table_index": row.get("table_index"),
        "source_row_number": row["source_row_number"],
        "entity_kind": row["entity_kind"],
        "values": row["values"],
        "normalized_values": row["normalized_values"],
        "issues": row["issues"],
        "status": row["status"],
        "action": row.get("action"),
        "confidence": float(row.get("confidence") or 0),
        "evidence": row.get("evidence") or [],
    }


def _write_meta_sheet(workbook: Workbook, *, source_standard: dict | None) -> None:
    meta_sheet = workbook.create_sheet(META_SHEET)
    meta_sheet.append(["key", "value"])
    meta_sheet.append(["schema_version", SCHEMA_VERSION])
    meta_sheet.append(["exported_at", datetime.now(timezone.utc).isoformat()])
    meta_sheet.append(["source_standard_id", source_standard.get("id") if source_standard else ""])
    meta_sheet.append(["source_standard_code", source_standard.get("code") if source_standard else ""])
    meta_sheet.sheet_state = "hidden"


def _create_sheet(workbook: Workbook, title: str, headers: list[str], rows: list[list[Any]]) -> None:
    sheet = workbook.create_sheet(title)
    sheet.append(headers)
    for cell in sheet[1]:
        cell.font = Font(bold=True)
    for row in rows:
        sheet.append(row)
    sheet.freeze_panes = "A2"
    for index, _header in enumerate(headers, start=1):
        sheet.column_dimensions[chr(64 + index)].width = 24


def _save_workbook(workbook: Workbook) -> bytes:
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _json_text(value: Any) -> str:
    return json.dumps(value or {}, ensure_ascii=False, sort_keys=True)


def _parse_json_object(value: Any, field: str) -> tuple[dict, dict | None]:
    text = _trim_text(value)
    if not text:
        return {}, None
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError:
        return {}, _issue("invalid_json", field, "字段必须是合法 JSON 对象")
    if not isinstance(parsed, dict):
        return {}, _issue("invalid_json", field, "字段必须是 JSON 对象")
    return parsed, None


def _parse_list(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    text = _trim_text(value)
    if not text:
        return []
    if text.startswith("["):
        try:
            parsed = json.loads(text)
        except json.JSONDecodeError:
            parsed = None
        if isinstance(parsed, list):
            return [str(item).strip() for item in parsed if str(item).strip()]
    return [item.strip() for item in re.split(r"[,;；，]", text) if item.strip()]


def _parse_bool(value: Any) -> bool | None:
    if isinstance(value, bool):
        return value
    text = _trim_text(value).lower()
    if not text:
        return False
    if text in {"true", "1", "yes", "y", "是", "必填"}:
        return True
    if text in {"false", "0", "no", "n", "否", "可选"}:
        return False
    return None


def _normalize_status(value: Any, allowed: set[str]) -> str | None:
    text = _trim_text(value).lower()
    if not text:
        return "active"
    return text if text in allowed else None


def _normalize_code(value: Any) -> str:
    return _trim_text(value)


def _code_key(value: Any) -> str:
    return _normalize_code(value).lower()


def _normalize_header(value: Any) -> str:
    return _trim_text(value).lower()


def _trim_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _blank_to_none(value: Any) -> str | None:
    text = _trim_text(value)
    return text or None


def _int_or_none(value: Any) -> int | None:
    if value is None or _trim_text(value) == "":
        return None
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _cell_value(value: Any) -> Any:
    if isinstance(value, datetime):
        return value.isoformat()
    return value


def _row_is_empty(values: dict[str, Any]) -> bool:
    return all(_trim_text(value) == "" for value in values.values())


def _issue(code: str, field: str, message: str, *, severity: Literal["error", "warning"] = "error") -> dict:
    return {"code": code, "field": field, "message": message, "severity": severity}
